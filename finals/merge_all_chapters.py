#!/usr/bin/env python3.11
# -*- coding: utf-8 -*-
"""
17個の章ごとのWord文書を結合するスクリプト
"""

from docx import Document

def merge_all_chapters(files, output):
    """
    17個のWord文書を結合
    
    Args:
        files: 結合するファイルのリスト
        output: 出力ファイル名
    """
    # 最初のファイルを基準として読み込み
    merged_doc = Document(files[0])
    
    # 2つ目以降のファイルを追加
    for file in files[1:]:
        # ページ区切りを追加
        merged_doc.add_page_break()
        
        # ファイルを読み込み
        doc = Document(file)
        
        # すべての段落をコピー
        for para in doc.paragraphs:
            new_para = merged_doc.add_paragraph(para.text, style=para.style)
            # フォーマットをコピー
            new_para.paragraph_format.alignment = para.paragraph_format.alignment
            new_para.paragraph_format.left_indent = para.paragraph_format.left_indent
            new_para.paragraph_format.right_indent = para.paragraph_format.right_indent
            new_para.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
            new_para.paragraph_format.space_before = para.paragraph_format.space_before
            new_para.paragraph_format.space_after = para.paragraph_format.space_after
            new_para.paragraph_format.line_spacing = para.paragraph_format.line_spacing
        
        # すべてのテーブルをコピー
        for table in doc.tables:
            # テーブルの行数と列数を取得
            rows = len(table.rows)
            cols = len(table.columns)
            
            # 新しいテーブルを作成
            new_table = merged_doc.add_table(rows=rows, cols=cols)
            new_table.style = table.style
            
            # セルの内容をコピー
            for i in range(rows):
                for j in range(cols):
                    try:
                        new_table.cell(i, j).text = table.cell(i, j).text
                    except:
                        pass  # セルが結合されている場合はスキップ
    
    # 保存
    merged_doc.save(output)
    print(f"結合完了: {output}")

if __name__ == "__main__":
    files = [
        "quickstart.docx",
        "preface.docx",
        "chapter1_final.docx",
        "chapter2_final.docx",
        "chapter3_final.docx",
        "chapter4_final.docx",
        "chapter5_final.docx",
        "chapter6_final.docx",
        "chapter7_final.docx",
        "chapter8_final.docx",
        "chapter9_final.docx",
        "chapter10_final.docx",
        "chapter11_final.docx",
        "chapter12_final.docx",
        "afterword.docx",
        "appendix_b_final.docx",
        "appendix_c_final.docx"
    ]
    
    output = "complete_manuscript_v21_final.docx"
    
    merge_all_chapters(files, output)
