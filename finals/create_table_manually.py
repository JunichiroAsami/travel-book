#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# chapter3のWord文書を読み込み
doc = Document('chapter3_heading1_test.docx')

# テーブルデータ
table_data = [
    ['ホテル名', '予算', '朝食', 'プール', 'ルーフトップバー', 'トゥクトゥク乗車'],
    ['Avani+ Riverside', '◎ 約1万円台', '◎ 豊富＋景観', '◎ インフィニティプール', '◎ プール併設バー', '○ 川沿い移動可'],
    ['Eastin Grand Sathorn', '◎ リーズナブル', '◎ フルーツ・オムレツ', '◎ 屋外プール', '○ 屋外バーあり', '○ 近隣で利用可'],
    ['Marriott Surawongse', '○ プロモ価格で可', '◎ 高評価ビュッフェ', '◎ インフィニティプール', '◎ 33階ルーフトップバー', '○ シーロム界隈で可'],
    ['Sivatel バンコク', '◎ コスパ高評価', '○ 種類少なめも質あり', '◎ 静かな屋外プール', '× 記載なし', '○ BTS周辺で可'],
]

# 「おすすめホテル」の見出しを探して、その後にテーブルを挿入
found_heading = False
insert_index = None

for i, para in enumerate(doc.paragraphs):
    if 'おすすめホテル' in para.text and para.style.name == 'Heading 1':
        found_heading = True
        # 見出しの後の段落を探す（最後のホテルの説明の後）
        # 「トゥクトゥクも楽しみたい方へ」の見出しの直前に挿入
        continue
    
    if found_heading and 'トゥクトゥクも楽しみたい方へ' in para.text:
        insert_index = i
        break

if insert_index is None:
    print("⚠️ 挿入位置が見つかりませんでした")
    exit(1)

print(f"✅ 挿入位置を発見: 段落{insert_index}の前")

# テーブルを作成
table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
table.style = 'Table Grid'

# テーブルにデータを入力
for row_idx, row_data in enumerate(table_data):
    for col_idx, cell_data in enumerate(row_data):
        cell = table.rows[row_idx].cells[col_idx]
        cell.text = cell_data
        
        # ヘッダー行は太字
        if row_idx == 0:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

# テーブルを適切な位置に移動
# （python-docxでは直接移動できないため、新しい文書を作成）
new_doc = Document()

for i, para in enumerate(doc.paragraphs):
    if i == insert_index:
        # テーブルを挿入
        new_table = new_doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        new_table.style = 'Table Grid'
        
        for row_idx, row_data in enumerate(table_data):
            for col_idx, cell_data in enumerate(row_data):
                cell = new_table.rows[row_idx].cells[col_idx]
                cell.text = cell_data
                
                if row_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
        
        # 空行を追加
        new_doc.add_paragraph()
    
    # 元の段落をコピー
    new_para = new_doc.add_paragraph(para.text, style=para.style)

# 元のテーブルもコピー
for table in doc.tables:
    new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
    new_table.style = table.style
    
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            new_table.rows[row_idx].cells[col_idx].text = cell.text

# 保存
new_doc.save('chapter3_with_manual_table.docx')

print("✅ テーブルを手動で作成しました: chapter3_with_manual_table.docx")
