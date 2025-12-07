#!/usr/bin/env python3.11
"""
v19のフォーマット（目次、ヘッダー、ページ番号）を使用してv21を作成
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy

print("=" * 80)
print("v19のフォーマットでv21を作成")
print("=" * 80)
print()

# v19をテンプレートとして読み込み
template = Document("complete_manuscript_v19_final.docx")
print("✓ v19をテンプレートとして読み込みました")

# ヘッダーとフッターを保存
header = template.sections[0].header
footer = template.sections[0].footer
page_setup = template.sections[0]

print(f"✓ ヘッダー: {len(header.paragraphs)}個の段落")
print(f"✓ フッター: {len(footer.paragraphs)}個の段落")

# v21の内容を読み込み（17個の個別章ファイルから）
chapter_files = [
    "quickstart.md",
    "preface.md",
    "chapter1_final.md",
    "chapter2_final.md",
    "chapter3_final.md",
    "chapter4_final.md",
    "chapter5_final.md",
    "chapter6_final.md",
    "chapter7_final.md",
    "chapter8_final.md",
    "chapter9_final.md",
    "chapter10_final.md",
    "chapter11_final.md",
    "chapter12_final.md",
    "afterword.md",
    "appendix_b_final.md",
    "appendix_c_final.md",
]

# v19の本文をすべて削除（目次以降）
# 目次は最初の数段落に含まれているので、それを保持
toc_end_index = 0
for i, para in enumerate(template.paragraphs):
    if para.style.name == 'Heading 1':
        toc_end_index = i
        break

print(f"✓ 目次は最初の{toc_end_index}段落")

# すべての段落を削除（目次は後で再生成）
paragraphs_to_remove = list(template.paragraphs)
for para in paragraphs_to_remove:
    p = para._element
    p.getparent().remove(p)

print(f"✓ {len(paragraphs_to_remove)}個の段落を削除しました")

# 目次を追加
print()
print("目次を追加中...")

# 目次タイトル
toc_title = template.add_paragraph()
toc_title_run = toc_title.add_run("TABLE OF CONTENTS")
toc_title_run.font.size = Pt(16)
toc_title_run.bold = True
toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 背景色を青に設定（v19と同じ）
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

shading_elm = OxmlElement('w:shd')
shading_elm.set(qn('w:fill'), '00B0F0')  # 青色
toc_title._element.get_or_add_pPr().append(shading_elm)

print("✓ 目次タイトルを追加しました")

# v21の内容を追加（Pandocで変換したファイルから）
print()
print("v21の内容を追加中...")

for i, chapter_file in enumerate(chapter_files):
    chapter_docx = f"chapter_{i:02d}.docx"
    chapter_doc = Document(chapter_docx)
    
    # ページ区切りを追加
    template.add_page_break()
    
    # 段落をコピー
    for element in chapter_doc.element.body:
        template.element.body.append(deepcopy(element))
    
    print(f"  ✓ {chapter_file} を追加しました")

# 保存
output_file = "complete_manuscript_v21_final.docx"
template.save(output_file)

print()
print("=" * 80)
print(f"✓ 完了: {output_file}")
print("=" * 80)

# 検証
doc = Document(output_file)
heading1_count = sum(1 for para in doc.paragraphs if para.style.name == 'Heading 1')
heading2_count = sum(1 for para in doc.paragraphs if para.style.name == 'Heading 2')
heading3_count = sum(1 for para in doc.paragraphs if para.style.name == 'Heading 3')

print()
print("【検証結果】")
print(f"  見出し1の数: {heading1_count}")
print(f"  見出し2の数: {heading2_count}")
print(f"  見出し3の数: {heading3_count}")
print(f"  総段落数: {len(doc.paragraphs)}")
print(f"  セクション数: {len(doc.sections)}")
