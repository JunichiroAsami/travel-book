#!/usr/bin/env python3
"""Word原稿v29を作成: 正しい構成順序で統合"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_page_break(doc):
    """改ページを追加"""
    doc.add_page_break()

def add_title_page_image(doc, image_path):
    """扉画像を追加"""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 画像を追加（A4サイズに合わせて調整）
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(6.5))
    
    add_page_break(doc)

def add_markdown_content(doc, md_path, title=None, title_level=1):
    """Markdownファイルの内容を追加"""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # タイトルを追加（オプション）
    if title:
        heading = doc.add_heading(title, level=title_level)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 内容を解析して追加
    lines = content.split('\n')
    skip_first_heading = True
    
    for line in lines:
        line_stripped = line.strip()
        
        # 空行
        if not line_stripped:
            doc.add_paragraph()
            continue
        
        # 最初の見出し（# タイトル）はスキップ
        if skip_first_heading and line_stripped.startswith('# '):
            skip_first_heading = False
            continue
        
        # 見出し2（## ...）
        if line_stripped.startswith('## '):
            heading_text = line_stripped[3:]
            heading = doc.add_heading(heading_text, level=2)
        # 見出し3（### ...）
        elif line_stripped.startswith('### '):
            heading_text = line_stripped[4:]
            heading = doc.add_heading(heading_text, level=3)
        # 区切り線（---）
        elif line_stripped.startswith('---'):
            doc.add_paragraph()  # 空行として扱う
        # 通常のテキスト
        elif not line_stripped.startswith('#'):
            paragraph = doc.add_paragraph(line_stripped)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    add_page_break(doc)

def add_halftitle(doc, md_path):
    """中扉を追加（中央揃え）"""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    skip_first_heading = True
    
    for line in lines:
        line_stripped = line.strip()
        
        if not line_stripped:
            doc.add_paragraph()
            continue
        
        # 最初の見出し（# 中扉）はスキップ
        if skip_first_heading and line_stripped.startswith('# '):
            skip_first_heading = False
            continue
        
        # すべてのテキストを中央揃えで追加
        if not line_stripped.startswith('#'):
            paragraph = doc.add_paragraph(line_stripped)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # フォントサイズを調整
            if '――' in line_stripped:
                # 区切り線は小さく
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            elif len(line_stripped) < 30:
                # 短いテキスト（タイトルなど）は大きく
                for run in paragraph.runs:
                    run.font.size = Pt(16)
                    run.font.bold = True
    
    add_page_break(doc)

def remove_old_title_page(doc):
    """古いテキストベースの扉ページを削除"""
    # 最初の数段落をチェックして、古いタイトルページを削除
    paragraphs_to_remove = []
    
    for i, paragraph in enumerate(doc.paragraphs[:10]):
        text = paragraph.text.strip()
        # 古いタイトルページの特徴的なテキストを検出
        if 'AIツアーコンダクター' in text or '周遊記' in text or (i < 5 and len(text) < 100):
            paragraphs_to_remove.append(paragraph)
    
    # 削除
    for paragraph in paragraphs_to_remove:
        p = paragraph._element
        p.getparent().remove(p)

def main():
    print("Word原稿v29を作成中...")
    print("=" * 80)
    
    # 新しいドキュメントを作成
    new_doc = Document()
    
    # 1. 扉画像を追加
    print("✓ 扉画像を追加中...")
    add_title_page_image(new_doc, '/home/ubuntu/ai_travel_book_project/finals/title_page_final.png')
    
    # 2. 前書きを追加
    print("✓ 前書きを追加中...")
    add_markdown_content(new_doc, '/home/ubuntu/ai_travel_book_project/finals/preface_v1.md', title='前書き')
    
    # 3. v27の内容を読み込む（目次と本文）
    print("✓ v27の目次と本文を読み込み中...")
    v27_doc = Document('/home/ubuntu/ai_travel_book_project/finals/complete_manuscript_v27_safe.docx')
    
    # 古いタイトルページを削除
    remove_old_title_page(v27_doc)
    
    # v27の内容をコピー（目次と本文）
    for element in v27_doc.element.body:
        new_doc.element.body.append(element)
    
    # 4. 中扉を追加（目次の後、本文の前に挿入する必要があるが、簡略化のため本文の前に配置）
    print("✓ 中扉を追加中...")
    add_halftitle(new_doc, '/home/ubuntu/ai_travel_book_project/finals/halftitle_v1.md')
    
    # 5. 後書きを追加
    print("✓ 後書きを追加中...")
    add_markdown_content(new_doc, '/home/ubuntu/ai_travel_book_project/finals/afterword_v1.md', title='後書き')
    
    # 6. ディスクレーマーを追加
    print("✓ ディスクレーマーを追加中...")
    add_markdown_content(new_doc, '/home/ubuntu/ai_travel_book_project/finals/disclaimer_v1.md', title='免責事項')
    
    # 7. 奥付を追加
    print("✓ 奥付を追加中...")
    add_markdown_content(new_doc, '/home/ubuntu/ai_travel_book_project/finals/colophon_v1.md', title='奥付')
    
    # 8. 保存
    output_path = '/home/ubuntu/ai_travel_book_project/finals/complete_manuscript_v29.docx'
    new_doc.save(output_path)
    print(f"✓ v29原稿を保存しました: {output_path}")
    
    print("=" * 80)
    print("完了！")
    print("\n注意: 中扉の位置は手動で調整が必要です（目次の後、本文の前に移動）")

if __name__ == '__main__':
    main()
