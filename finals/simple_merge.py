#!/usr/bin/env python3.11
# -*- coding: utf-8 -*-
"""
3つのWord文書をシンプルに結合するスクリプト
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_page_break(paragraph):
    """段落にページ区切りを追加"""
    run = paragraph.add_run()
    element = run._element
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    element.append(br)

def simple_merge(files, output):
    """
    Word文書をシンプルに結合
    
    Args:
        files: 結合するファイルのリスト
        output: 出力ファイル名
    """
    # 新しい空のドキュメントを作成
    merged_doc = Document()
    
    for i, file in enumerate(files):
        print(f"結合中: {file}")
        doc = Document(file)
        
        # ページ区切りを追加（最初のファイル以外）
        if i > 0:
            p = merged_doc.add_paragraph()
            add_page_break(p)
        
        # すべての要素をコピー
        for element in doc.element.body:
            merged_doc.element.body.append(element)
    
    # 保存
    merged_doc.save(output)
    print(f"結合完了: {output}")

if __name__ == "__main__":
    files = [
        "temp_part1.docx",
        "temp_part2.docx",
        "temp_part3.docx"
    ]
    
    output = "complete_manuscript_v21_final.docx"
    
    simple_merge(files, output)
