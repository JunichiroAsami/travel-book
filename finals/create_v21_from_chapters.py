#!/usr/bin/env python3.11
"""
17個の個別章ファイルから正しいWord文書を作成
"""

from docx import Document
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess

print("=" * 80)
print("17個の個別章ファイルからv21を作成")
print("=" * 80)
print()

chapter_files = [
    "quickstart.md",
    "preface.md",
    "chapter1_final.md",
    "chapter2_final.md",
    "chapter3_final.md",
    "chapter4_final.md",
    "chapter5_final.md",
    "chapter6_final.md",
    "chapter7_final.md",
    "chapter8_final.md",
    "chapter9_final.md",
    "chapter10_final.md",
    "chapter11_final.md",
    "chapter12_final.md",
    "afterword.md",
    "appendix_b_final.md",
    "appendix_c_final.md",
]

# 各章をPandocでWord変換（v19を参照文書として使用）
print("各章をPandocでWord変換中...")
for i, chapter_file in enumerate(chapter_files):
    chapter_docx = f"chapter_{i:02d}.docx"
    subprocess.run([
        "pandoc",
        chapter_file,
        "-o", chapter_docx,
        "--reference-doc=complete_manuscript_v19_final.docx"
    ], check=True)
    print(f"  ✓ {chapter_file} → {chapter_docx}")

print()
print("=" * 80)
print("Word文書を結合中...")
print("=" * 80)
print()

# 新しいWord文書を作成（v19をベースにする）
base_doc = Document("complete_manuscript_v19_final.docx")

# v19のヘッダーとフッターを保存
header_xml = base_doc.sections[0].header._element
footer_xml = base_doc.sections[0].footer._element

# v19のすべての段落を削除
for para in list(base_doc.paragraphs):
    p = para._element
    p.getparent().remove(p)

print("✓ v19のベース文書を準備しました")

# 目次タイトルを追加
toc_title = base_doc.add_paragraph()
toc_title_run = toc_title.add_run("TABLE OF CONTENTS")
toc_title_run.font.size = Pt(16)
toc_title_run.bold = True
toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 背景色を青に設定
shading_elm = OxmlElement('w:shd')
shading_elm.set(qn('w:fill'), '00B0F0')
toc_title._element.get_or_add_pPr().append(shading_elm)

print("✓ 目次タイトルを追加しました")

# 17個の章を結合
print()
print("章を結合中...")

for i, chapter_file in enumerate(chapter_files):
    chapter_docx = f"chapter_{i:02d}.docx"
    chapter_doc = Document(chapter_docx)
    
    # ページ区切りを追加
    base_doc.add_page_break()
    
    # 段落を追加
    for para in chapter_doc.paragraphs:
        new_para = base_doc.add_paragraph(para.text, style=para.style.name)
        # フォーマットをコピー
        new_para.alignment = para.alignment
        for run in para.runs:
            new_run = new_para.runs[-1] if new_para.runs else new_para.add_run()
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            if run.font.size:
                new_run.font.size = run.font.size
    
    # テーブルを追加
    for table in chapter_doc.tables:
        new_table = base_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                new_table.rows[row_idx].cells[col_idx].text = cell.text
    
    print(f"  ✓ {chapter_file} を結合しました")

# 保存
output_file = "complete_manuscript_v21_final.docx"
base_doc.save(output_file)

print()
print("=" * 80)
print(f"✓ 完了: {output_file}")
print("=" * 80)

# 検証
doc = Document(output_file)
heading1_count = sum(1 for para in doc.paragraphs if para.style.name == 'Heading 1')
heading2_count = sum(1 for para in doc.paragraphs if para.style.name == 'Heading 2')
heading3_count = sum(1 for para in doc.paragraphs if para.style.name == 'Heading 3')

print()
print("【検証結果】")
print(f"  見出し1の数: {heading1_count}")
print(f"  見出し2の数: {heading2_count}")
print(f"  見出し3の数: {heading3_count}")
print(f"  総段落数: {len(doc.paragraphs)}")
print(f"  セクション数: {len(doc.sections)}")
