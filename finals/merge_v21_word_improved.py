#!/usr/bin/env python3.11
# -*- coding: utf-8 -*-
"""
v21の3つのWord文書を正しく結合するスクリプト（改良版）
"""

from docx import Document
from copy import deepcopy

def merge_word_documents_improved(files, output):
    """
    複数のWord文書を結合（見出しレベルを保持）
    
    Args:
        files: 結合するファイルのリスト
        output: 出力ファイル名
    """
    # 最初のファイルを基準として読み込み
    merged_doc = Document(files[0])
    
    # 2つ目以降のファイルを追加
    for file in files[1:]:
        # ページ区切りを追加
        merged_doc.add_page_break()
        
        # ファイルを読み込み
        doc = Document(file)
        
        # すべての段落をコピー
        for para in doc.paragraphs:
            new_para = merged_doc.add_paragraph(para.text, style=para.style)
            # フォーマットをコピー
            new_para.paragraph_format.alignment = para.paragraph_format.alignment
            new_para.paragraph_format.left_indent = para.paragraph_format.left_indent
            new_para.paragraph_format.right_indent = para.paragraph_format.right_indent
            new_para.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
            new_para.paragraph_format.space_before = para.paragraph_format.space_before
            new_para.paragraph_format.space_after = para.paragraph_format.space_after
            new_para.paragraph_format.line_spacing = para.paragraph_format.line_spacing
        
        # すべてのテーブルをコピー
        for table in doc.tables:
            # テーブルの行数と列数を取得
            rows = len(table.rows)
            cols = len(table.columns)
            
            # 新しいテーブルを作成
            new_table = merged_doc.add_table(rows=rows, cols=cols)
            new_table.style = table.style
            
            # セルの内容をコピー
            for i in range(rows):
                for j in range(cols):
                    try:
                        new_table.cell(i, j).text = table.cell(i, j).text
                    except:
                        pass  # セルが結合されている場合はスキップ
    
    # 保存
    merged_doc.save(output)
    print(f"結合完了: {output}")

if __name__ == "__main__":
    files = [
        "v21_part1.docx",
        "v21_part2.docx",
        "v21_part3.docx"
    ]
    
    output = "complete_manuscript_v21_final.docx"
    
    merge_word_documents_improved(files, output)
