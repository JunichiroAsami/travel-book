#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# chapter3のWord文書を読み込み
doc = Document('chapter3_no_table_v2.docx')

# テーブルデータ
table_data = [
    ['ホテル名', '予算', '朝食', 'プール', 'ルーフトップバー', 'トゥクトゥク乗車'],
    ['Avani+ Riverside', '◎ 約1万円台', '◎ 豊富＋景観', '◎ インフィニティプール', '◎ プール併設バー', '○ 川沿い移動可'],
    ['Eastin Grand Sathorn', '◎ リーズナブル', '◎ フルーツ・オムレツ', '◎ 屋外プール', '○ 屋外バーあり', '○ 近隣で利用可'],
    ['Marriott Surawongse', '○ プロモ価格で可', '◎ 高評価ビュッフェ', '◎ インフィニティプール', '◎ 33階ルーフトップバー', '○ シーロム界隈で可'],
    ['Sivatel バンコク', '◎ コスパ高評価', '○ 種類少なめも質あり', '◎ 静かな屋外プール', '× 記載なし', '○ BTS周辺で可'],
]

# TABLE_PLACEHOLDERを探す
placeholder_index = None

for i, para in enumerate(doc.paragraphs):
    if 'TABLE_PLACEHOLDER' in para.text:
        placeholder_index = i
        print(f"✅ プレースホルダーを発見: 段落{i}")
        break

if placeholder_index is None:
    print("⚠️ プレースホルダーが見つかりませんでした")
    exit(1)

# プレースホルダーの段落を削除
placeholder_para = doc.paragraphs[placeholder_index]
p = placeholder_para._element
p.getparent().remove(p)

# 新しい文書を作成して、テーブルを適切な位置に挿入
new_doc = Document()

# 元の文書の内容をコピー
for i, para in enumerate(doc.paragraphs):
    if i == placeholder_index:
        # テーブルを挿入
        table = new_doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        table.style = 'Table Grid'
        
        for row_idx, row_data in enumerate(table_data):
            for col_idx, cell_data in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = cell_data
                
                # ヘッダー行は太字
                if row_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
        
        # 空行を追加
        new_doc.add_paragraph()
    
    # 元の段落をコピー
    new_para = new_doc.add_paragraph(para.text, style=para.style)

# 元のテーブルもコピー
for table in doc.tables:
    if table.rows and table.rows[0].cells:
        col_count = len(table.rows[0].cells)
        new_table = new_doc.add_table(rows=len(table.rows), cols=col_count)
        new_table.style = table.style
        
        for row_idx, row in enumerate(table.rows):
            for col_idx in range(min(col_count, len(row.cells))):
                new_table.rows[row_idx].cells[col_idx].text = row.cells[col_idx].text

# 保存
new_doc.save('chapter3_with_table.docx')

print("✅ テーブルを挿入しました: chapter3_with_table.docx")

# 検証
doc_verify = Document('chapter3_with_table.docx')
print(f"\n検証: 総テーブル数 = {len(doc_verify.tables)}")

for i, table in enumerate(doc_verify.tables, 1):
    if table.rows:
        col_count = len(table.rows[0].cells)
        print(f"  テーブル{i}: {len(table.rows)}行 × {col_count}列")
