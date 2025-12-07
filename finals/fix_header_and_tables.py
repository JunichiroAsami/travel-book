#!/usr/bin/env python3
"""
v26原稿のページ番号を上部に移動し、テーブルの枠線を太く・黒くするスクリプト
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

def add_page_number(paragraph):
    """段落にページ番号を追加"""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    
    return run

def set_table_borders(table, width=12, color='000000'):
    """テーブルの枠線を設定"""
    tbl = table._element
    tblPr = tbl.tblPr
    
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # 既存のtblBordersを削除
    for borders in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(borders)
    
    # 新しいtblBordersを作成
    tblBorders = OxmlElement('w:tblBorders')
    
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(width))  # 1/8ポイント単位（12 = 1.5pt）
        border.set(qn('w:color'), color)
        tblBorders.append(border)
    
    tblPr.append(tblBorders)

INPUT_DOCX = 'complete_manuscript_v26_final_fixed.docx'
OUTPUT_DOCX = 'complete_manuscript_v26_final_v2.docx'

print("=" * 60)
print("ヘッダー・テーブル修正スクリプト")
print("=" * 60)

doc = Document(INPUT_DOCX)

# ========================================
# 1. ヘッダー・フッターの修正
# ========================================

print("\n【1. ヘッダー・フッターの修正】")

section = doc.sections[0]

# ヘッダー: テーブルを使って左側に書籍タイトル、右側にページ番号
header = section.header
for paragraph in header.paragraphs:
    paragraph.clear()

# ヘッダーにテーブルを追加
from docx.shared import Inches
header_table = header.add_table(rows=1, cols=2, width=Inches(6))
header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 列幅を設定（左側を広く、右側を狭く）
header_table.columns[0].width = Pt(400)
header_table.columns[1].width = Pt(80)

# 左側のセル: 書籍タイトル
left_cell = header_table.rows[0].cells[0]
left_para = left_cell.paragraphs[0]
left_para.text = "AIツアーコンダクターと歩く ベトナム・タイ周遊記"
left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
left_para.runs[0].font.size = Pt(10)

# 右側のセル: ページ番号（青い背景）
right_cell = header_table.rows[0].cells[1]
right_para = right_cell.paragraphs[0]
right_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_page_number(right_para)
right_para.runs[0].font.size = Pt(10)
right_para.runs[0].font.color.rgb = RGBColor(255, 255, 255)  # 白文字

# セルの背景色を青に設定
shading = OxmlElement('w:shd')
shading.set(qn('w:fill'), '7CB8C5')  # 青色
right_cell._element.get_or_add_tcPr().append(shading)

# テーブルの枠線を削除
set_table_borders(header_table, width=0, color='FFFFFF')

print("  ✅ ヘッダーを修正（左: タイトル、右: ページ番号）")

# フッターをクリア
footer = section.footer
for paragraph in footer.paragraphs:
    paragraph.clear()

print("  ✅ フッターをクリア")

# 表紙ページの設定
section.different_first_page_header_footer = True

first_page_header = section.first_page_header
first_page_footer = section.first_page_footer

for paragraph in first_page_header.paragraphs:
    paragraph.clear()

for paragraph in first_page_footer.paragraphs:
    paragraph.clear()

print("  ✅ 表紙ページのヘッダー・フッターをクリア")

# ========================================
# 2. テーブルの枠線を太く・黒くする
# ========================================

print("\n【2. テーブルの枠線を修正】")

table_count = 0
for table in doc.tables:
    set_table_borders(table, width=12, color='000000')  # 1.5pt、黒
    table_count += 1

print(f"  ✅ {table_count}個のテーブルの枠線を太く・黒くしました")

# ========================================
# 3. 保存
# ========================================

doc.save(OUTPUT_DOCX)

print(f"\n✅ 完了: {OUTPUT_DOCX}")
print("=" * 60)
