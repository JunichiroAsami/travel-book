#!/usr/bin/env python3.11
"""
目次ページを作成してv21に追加
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

print("=" * 80)
print("目次ページを作成してv21に追加")
print("=" * 80)
print()

# v21を読み込み
doc = Document("complete_manuscript_v21_final.docx")
print(f"✓ v21を読み込みました（{len(doc.paragraphs)}段落）")

# 見出し1と見出し2を抽出
headings = []
for para in doc.paragraphs:
    if para.style.name == 'Heading 1':
        headings.append((1, para.text))
    elif para.style.name == 'Heading 2':
        headings.append((2, para.text))

print(f"✓ 見出しを抽出しました（見出し1: {sum(1 for h in headings if h[0] == 1)}個、見出し2: {sum(1 for h in headings if h[0] == 2)}個）")

# 新しいWord文書を作成（目次用）
toc_doc = Document()

# v21のスタイルとページ設定をコピー
for section in doc.sections:
    toc_section = toc_doc.sections[0]
    toc_section.page_width = section.page_width
    toc_section.page_height = section.page_height
    toc_section.left_margin = section.left_margin
    toc_section.right_margin = section.right_margin
    toc_section.top_margin = section.top_margin
    toc_section.bottom_margin = section.bottom_margin
    break

print("✓ ページ設定をコピーしました")

# 目次タイトルを追加
toc_title = toc_doc.add_paragraph()
toc_title_run = toc_title.add_run("TABLE OF CONTENTS")
toc_title_run.font.size = Pt(16)
toc_title_run.bold = True
toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 背景色を青に設定
shading_elm = OxmlElement('w:shd')
shading_elm.set(qn('w:fill'), '00B0F0')
toc_title._element.get_or_add_pPr().append(shading_elm)

print("✓ 目次タイトルを追加しました")

# 見出しを目次として追加（見出し2まで）
for level, text in headings[:50]:  # 最初の50個のみ（目次が長くなりすぎないように）
    toc_entry = toc_doc.add_paragraph()
    if level == 1:
        toc_entry_run = toc_entry.add_run(text)
        toc_entry_run.bold = True
        toc_entry_run.font.size = Pt(11)
    else:
        toc_entry_run = toc_entry.add_run(f"  {text}")
        toc_entry_run.font.size = Pt(10)

print(f"✓ 目次エントリを追加しました（{min(50, len(headings))}個）")

# ページ区切りを追加
toc_doc.add_page_break()

# v21の内容を追加
for element in doc.element.body:
    toc_doc.element.body.append(element)

print("✓ v21の内容を追加しました")

# 保存
output_file = "complete_manuscript_v21_final_with_toc.docx"
toc_doc.save(output_file)

print()
print("=" * 80)
print(f"✓ 完了: {output_file}")
print("=" * 80)
