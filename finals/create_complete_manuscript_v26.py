#!/usr/bin/env python3
"""
完全なWord原稿（v26）を作成する自動化スクリプト

処理フロー:
1. 全MDファイルを結合
2. 欠落テーブルの位置にプレースホルダーを挿入
3. Pandocで変換
4. python-docxで各プレースホルダーにテーブルを挿入
5. ヘッダー・フッター・ページ番号を追加
"""

import os
import re
import json
import subprocess
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# === 設定 ===
WORK_DIR = '/home/ubuntu/ai_travel_book_project/finals'
os.chdir(WORK_DIR)

# MDファイルのリスト（順序重要）
MD_FILES = [
    'v21_part_00_quickstart.md',
    'introduction_final.md',
    'chapter1_final.md',
    'chapter2_final.md',
    'chapter3_final.md',
    'chapter4_final.md',
    'chapter5_final.md',
    'chapter6_final.md',
    'chapter7_final.md',
    'chapter8_final.md',
    'chapter9_final.md',
    'chapter10_final.md',
    'chapter11_final.md',
    'chapter12_final.md',
    'conclusion_final.md',
    'appendix_a_final.md',
    'appendix_b_final.md',
    'appendix_c_final.md'
]

# 欠落テーブルの情報（ファイル名、行番号、プレースホルダーID）
MISSING_TABLES = [
    {'file': 'chapter3_final.md', 'line': 69, 'id': 'ch3_t1'},
    {'file': 'chapter3_final.md', 'line': 140, 'id': 'ch3_t2'},
    {'file': 'chapter3_final.md', 'line': 147, 'id': 'ch3_t3'},
    {'file': 'chapter3_final.md', 'line': 156, 'id': 'ch3_t4'},
    {'file': 'chapter3_final.md', 'line': 283, 'id': 'ch3_t5'},
    {'file': 'chapter3_final.md', 'line': 343, 'id': 'ch3_t6'},
    {'file': 'chapter4_final.md', 'line': 214, 'id': 'ch4_t1'},
    {'file': 'chapter4_final.md', 'line': 221, 'id': 'ch4_t2'},
    {'file': 'chapter4_final.md', 'line': 249, 'id': 'ch4_t3'},
    {'file': 'chapter4_final.md', 'line': 307, 'id': 'ch4_t4'}
]

COMBINED_MD = 'combined_manuscript_v26.md'
COMBINED_MD_WITH_PLACEHOLDERS = 'combined_manuscript_v26_with_placeholders.md'
DOCX_WITH_PLACEHOLDERS = 'manuscript_v26_with_placeholders.docx'
FINAL_DOCX = 'complete_manuscript_v26_final.docx'
TABLE_DATA_JSON = 'missing_tables_data.json'

print("=" * 60)
print("完全なWord原稿（v26）作成スクリプト")
print("=" * 60)

# === ステップ1: 全MDファイルを結合 ===
print("\n【ステップ1】全MDファイルを結合中...")

combined_content = []
for md_file in MD_FILES:
    if not os.path.exists(md_file):
        print(f"⚠️  警告: {md_file} が見つかりません。スキップします。")
        continue
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
        combined_content.append(content)
        print(f"  ✅ {md_file}")

# 結合したMDファイルを保存
with open(COMBINED_MD, 'w', encoding='utf-8') as f:
    f.write('\n\n'.join(combined_content))

print(f"✅ 結合完了: {COMBINED_MD}")

# === ステップ2: 欠落テーブルの位置にプレースホルダーを挿入 ===
print("\n【ステップ2】欠落テーブルにプレースホルダーを挿入中...")

# 各ファイルごとに処理
file_contents = {}
for md_file in MD_FILES:
    if os.path.exists(md_file):
        with open(md_file, 'r', encoding='utf-8') as f:
            file_contents[md_file] = f.read()

# 欠落テーブルをプレースホルダーに置き換え
for table_info in MISSING_TABLES:
    filename = table_info['file']
    target_line = table_info['line']
    table_id = table_info['id']
    
    if filename not in file_contents:
        continue
    
    lines = file_contents[filename].split('\n')
    
    # テーブルの開始行と終了行を特定
    table_start = target_line - 1  # 0-indexed
    table_end = table_start
    
    for i in range(table_start, len(lines)):
        if re.match(r'^\|.+\|$', lines[i]):
            table_end = i
        else:
            break
    
    # テーブルをプレースホルダーに置き換え
    placeholder = f'【テーブル挿入位置:{table_id}】'
    lines[table_start:table_end+1] = [placeholder]
    
    file_contents[filename] = '\n'.join(lines)
    print(f"  ✅ {table_id}: {filename} 行{target_line}")

# プレースホルダー付きの結合MDファイルを作成
combined_with_placeholders = []
for md_file in MD_FILES:
    if md_file in file_contents:
        combined_with_placeholders.append(file_contents[md_file])

with open(COMBINED_MD_WITH_PLACEHOLDERS, 'w', encoding='utf-8') as f:
    f.write('\n\n'.join(combined_with_placeholders))

print(f"✅ プレースホルダー挿入完了: {COMBINED_MD_WITH_PLACEHOLDERS}")

# === ステップ3: Pandocで変換 ===
print("\n【ステップ3】Pandocで変換中...")

pandoc_cmd = [
    'pandoc',
    COMBINED_MD_WITH_PLACEHOLDERS,
    '-o', DOCX_WITH_PLACEHOLDERS,
    '--from=markdown',
    '--to=docx'
]

result = subprocess.run(pandoc_cmd, capture_output=True, text=True)
if result.returncode != 0:
    print(f"❌ Pandoc変換エラー: {result.stderr}")
    exit(1)

print(f"✅ Pandoc変換完了: {DOCX_WITH_PLACEHOLDERS}")

# === ステップ4: python-docxでテーブルを挿入 ===
print("\n【ステップ4】テーブルを手動で挿入中...")

# テーブルデータを読み込み
with open(TABLE_DATA_JSON, 'r', encoding='utf-8') as f:
    table_data_collection = json.load(f)

# Word文書を開く
doc = Document(DOCX_WITH_PLACEHOLDERS)

# 各プレースホルダーを検索してテーブルを挿入
for table_info in MISSING_TABLES:
    table_id = table_info['id']
    placeholder = f'【テーブル挿入位置:{table_id}】'
    
    if table_id not in table_data_collection:
        print(f"⚠️  警告: {table_id} のデータが見つかりません。")
        continue
    
    table_data = table_data_collection[table_id]['data']
    
    # プレースホルダーを検索
    found = False
    for para in doc.paragraphs:
        if placeholder in para.text:
            # テーブルを作成
            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            table.autofit = True
            
            # テーブルにデータを入力
            for i, row_data in enumerate(table_data):
                for j, cell_data in enumerate(row_data):
                    cell = table.rows[i].cells[j]
                    cell.text = cell_data
                    # フォントサイズを設定
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
            
            # プレースホルダーの段落をテーブルに置き換え
            parent = para._p.getparent()
            parent.insert(parent.index(para._p), table._element)
            parent.remove(para._p)
            
            print(f"  ✅ {table_id}: {len(table_data)}行×{len(table_data[0])}列のテーブルを挿入")
            found = True
            break
    
    if not found:
        print(f"  ⚠️  警告: {table_id} のプレースホルダーが見つかりません。")

# 一旦保存して再読み込み
temp_file = 'temp_with_tables.docx'
doc.save(temp_file)
print(f"\n✅ テーブル挿入完了。一旦保存します。")

# === ステップ4.5: 見出しレベルを修正 ===
# 文書を再読み込み
doc = Document(temp_file)
print("\n【ステップ4.5】見出しレベルを修正中...")

# Heading 2からHeading 1に昇格させる見出しのパターン
chapter_patterns = [
    r'^第\d+章：',  # 第4章〜第12章
    r'^おわりに$',
    r'^付録B：'
]

import re as regex_module

# 既存のHeading 1段落のスタイルオブジェクトを取得
h1_style = None
for para in doc.paragraphs:
    if para.style.name == 'Heading 1':
        h1_style = para.style
        break

if h1_style is None:
    print("⚠️  警告: Heading 1スタイルが見つかりません。")
else:
    for para in doc.paragraphs:
        if para.style.name == 'Heading 2':
            for pattern in chapter_patterns:
                if regex_module.match(pattern, para.text):
                    para.style = h1_style
                    print(f"  ✅ Heading 1に昇格: {para.text}")
                    break

# 文書を保存
doc.save(FINAL_DOCX)
print(f"\n✅ テーブル挿入完了: {FINAL_DOCX}")

# === ステップ5: 統計情報を表示 ===
print("\n【ステップ5】統計情報")

doc_final = Document(FINAL_DOCX)
print(f"  総段落数: {len(doc_final.paragraphs)}")
print(f"  総テーブル数: {len(doc_final.tables)}")

# Heading 1の数を数える
h1_count = sum(1 for p in doc_final.paragraphs if p.style.name == 'Heading 1')
print(f"  Heading 1の数: {h1_count}")

print("\n" + "=" * 60)
print("✅ 完全なWord原稿（v26）の作成が完了しました！")
print(f"出力ファイル: {FINAL_DOCX}")
print("=" * 60)
