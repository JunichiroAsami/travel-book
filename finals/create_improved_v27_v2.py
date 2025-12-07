#!/usr/bin/env python3
"""
改善版v27作成スクリプト v2

v26の手法（プレースホルダー付きMDファイル）を使用し、以下の改善を追加：
1. 本文レイアウト：段落の一字下げ、見出しスタイルの明確化
2. テーブル：枠線を2.0ptに
3. フォントと行間：統一感を持たせる
"""

import os
import json
import subprocess
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

print("="*60)
print("改善版v27作成スクリプト v2")
print("="*60)

# ステップ1: v26のプレースホルダー付きWord文書を基に作成
print("\nステップ1: v26のWord文書を読み込み")
doc = Document('complete_manuscript_v26_final.docx')
print(f"  ✅ 読み込み完了（段落数: {len(doc.paragraphs)}）")

# ステップ2: スタイルの改善
print("\nステップ2: スタイルの改善")

# 本文スタイルの設定
style_updated = 0
for style_name in ['Normal', 'Body Text', 'First Paragraph']:
    try:
        style = doc.styles[style_name]
        style.font.name = 'MS Mincho'  # 明朝体
        style.font.size = Pt(11)  # フォントサイズを11ptに
        if style.paragraph_format:
            style.paragraph_format.line_spacing = 1.5  # 行間を1.5倍に
            style.paragraph_format.first_line_indent = Inches(0.25)  # 一字下げ
            style.paragraph_format.space_after = Pt(6)  # 段落後のスペース
        style_updated += 1
    except:
        pass

# 見出しスタイルの設定
heading_configs = {
    'Heading 1': {'size': 18, 'bold': True, 'space_before': 24, 'space_after': 12},
    'Heading 2': {'size': 14, 'bold': True, 'space_before': 18, 'space_after': 6},
    'Heading 3': {'size': 12, 'bold': True, 'space_before': 12, 'space_after': 6},
}

for style_name, config in heading_configs.items():
    try:
        style = doc.styles[style_name]
        style.font.name = 'MS Gothic'  # ゴシック体
        style.font.size = Pt(config['size'])
        style.font.bold = config['bold']
        if style.paragraph_format:
            style.paragraph_format.space_before = Pt(config['space_before'])
            style.paragraph_format.space_after = Pt(config['space_after'])
            style.paragraph_format.first_line_indent = Inches(0)  # 見出しは字下げなし
        style_updated += 1
    except:
        pass

print(f"  ✅ {style_updated}個のスタイルを更新")

# 段落に一字下げを適用
paragraph_count = 0
for para in doc.paragraphs:
    if para.style.name in ['Normal', 'Body Text', 'First Paragraph', 'List Paragraph']:
        # 見出しやコードブロックは除外
        if not para.text.startswith('#') and 'Heading' not in para.style.name:
            try:
                if para.paragraph_format.first_line_indent is None or para.paragraph_format.first_line_indent == 0:
                    para.paragraph_format.first_line_indent = Inches(0.25)
                    para.paragraph_format.line_spacing = 1.5
                    paragraph_count += 1
            except:
                pass

print(f"  ✅ {paragraph_count}個の段落に一字下げと行間を適用")

# 保存
doc.save('temp_with_improved_styles_v27.docx')
print("\n✅ スタイル改善完了: temp_with_improved_styles_v27.docx")

# ステップ3: テーブルの枠線を太くする
print("\nステップ3: テーブルの枠線を2.0ptに変更")
doc = Document('temp_with_improved_styles_v27.docx')

table_count = 0
for table in doc.tables:
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # 既存のtblBordersを削除
    for child in list(tblPr):
        if child.tag == qn('w:tblBorders'):
            tblPr.remove(child)
    
    # 新しいtblBordersを追加
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '16')  # 2.0pt (8 * 2)
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)
    table_count += 1

print(f"  ✅ {table_count}個のテーブルの枠線を更新")

# 保存
doc.save('complete_manuscript_v27_final.docx')

print("\n" + "="*60)
print("✅ 改善版v27作成完了！")
print("="*60)
print(f"\n出力ファイル: complete_manuscript_v27_final.docx")
print(f"総段落数: {len(doc.paragraphs)}")
print(f"総テーブル数: {len(doc.tables)}")
print(f"総セクション数: {len(doc.sections)}")

# 統計情報
heading1_count = sum(1 for p in doc.paragraphs if p.style.name == 'Heading 1')
print(f"Heading 1の数: {heading1_count}")
