#!/usr/bin/env python3
"""
Word原稿にヘッダー・フッター・ページ番号を追加するスクリプト
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_page_number(paragraph):
    """段落にページ番号を追加"""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    
    return run

# === 設定 ===
INPUT_DOCX = 'complete_manuscript_v26_final.docx'
OUTPUT_DOCX = 'complete_manuscript_v26_with_headers.docx'
BOOK_TITLE = 'AIツアーコンダクターと歩く ベトナム・タイ周遊記'

print("=" * 60)
print("ヘッダー・フッター・ページ番号追加スクリプト")
print("=" * 60)

# Word文書を開く
doc = Document(INPUT_DOCX)

# === ヘッダーの設定 ===
print("\n【ステップ1】ヘッダーを設定中...")

# 最初のセクションのヘッダーを取得
section = doc.sections[0]
header = section.header

# ヘッダーの既存の段落をクリア
for paragraph in header.paragraphs:
    paragraph.clear()

# ヘッダーに書籍タイトルを追加
header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
header_para.text = BOOK_TITLE
header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
header_para.runs[0].font.size = Pt(10)
header_para.runs[0].font.italic = True

print(f"  ✅ ヘッダーに書籍タイトルを追加: {BOOK_TITLE}")

# === フッターの設定（ページ番号） ===
print("\n【ステップ2】フッターにページ番号を追加中...")

footer = section.footer

# フッターの既存の段落をクリア
for paragraph in footer.paragraphs:
    paragraph.clear()

# フッターにページ番号を追加
footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ページ番号を追加
add_page_number(footer_para)
footer_para.runs[0].font.size = Pt(10)

print("  ✅ フッターにページ番号を追加")

# === 表紙ページの設定（ページ番号なし） ===
print("\n【ステップ3】表紙ページの設定中...")

# 最初のセクションを「最初のページが異なる」に設定
section.different_first_page_header_footer = True

# 最初のページのヘッダーとフッターをクリア
first_page_header = section.first_page_header
first_page_footer = section.first_page_footer

for paragraph in first_page_header.paragraphs:
    paragraph.clear()

for paragraph in first_page_footer.paragraphs:
    paragraph.clear()

print("  ✅ 表紙ページのヘッダー・フッターをクリア")

# === 文書を保存 ===
doc.save(OUTPUT_DOCX)

print("\n" + "=" * 60)
print(f"✅ ヘッダー・フッター・ページ番号の追加が完了しました！")
print(f"出力ファイル: {OUTPUT_DOCX}")
print("=" * 60)
