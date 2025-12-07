#!/usr/bin/env python3
"""Word原稿v30を作成: 各要素を個別にWord化してマージ"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.shared import OxmlElement as OE

def set_font_meiryo(run, size_pt):
    """フォントをメイリオに設定"""
    run.font.name = 'メイリオ'
    run.font.size = Pt(size_pt)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), 'メイリオ')
    rFonts.set(qn('w:ascii'), 'メイリオ')
    rFonts.set(qn('w:hAnsi'), 'メイリオ')

def copy_paragraph(target_doc, source_para):
    """段落をコピー"""
    new_para = target_doc.add_paragraph()
    new_para.style = source_para.style
    new_para.alignment = source_para.alignment
    
    for run in source_para.runs:
        new_run = new_para.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        if run.font.size:
            new_run.font.size = run.font.size
        if run.font.name:
            new_run.font.name = run.font.name
        if run.font.color.rgb:
            new_run.font.color.rgb = run.font.color.rgb
    
    return new_para

def add_title_page(doc, image_path):
    """扉（画像）を追加"""
    print("  ✓ 扉（画像）を追加")
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(6))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

def add_markdown_section(doc, md_path, section_name):
    """Markdownファイルをセクションとして追加"""
    print(f"  ✓ {section_name}を追加")
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    for line in lines:
        line_stripped = line.strip()
        
        if not line_stripped:
            doc.add_paragraph()
            continue
        
        # 見出し
        if line_stripped.startswith('# '):
            heading_text = line_stripped.replace('# ', '')
            heading = doc.add_heading(heading_text, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue
        
        if line_stripped.startswith('## '):
            heading_text = line_stripped.replace('## ', '')
            heading = doc.add_heading(heading_text, level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue
        
        if line_stripped.startswith('### '):
            heading_text = line_stripped.replace('### ', '')
            heading = doc.add_heading(heading_text, level=3)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue
        
        # 通常のテキスト
        paragraph = doc.add_paragraph(line_stripped)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_disclaimer_page(doc, md_path):
    """免責事項ページを追加（野沢温泉形式）"""
    print("  ✓ 免責事項（野沢温泉形式）を追加")
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    skip_first_heading = True
    
    for line in lines:
        line_stripped = line.strip()
        
        if not line_stripped:
            doc.add_paragraph()
            continue
        
        if skip_first_heading and line_stripped.startswith('# '):
            skip_first_heading = False
            heading = doc.add_heading('免責事項（ディスクレーマー）', level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in heading.runs:
                set_font_meiryo(run, 12)
                run.font.bold = True
            continue
        
        if line_stripped and line_stripped[0].isdigit() and '. **' in line_stripped:
            parts = line_stripped.split('. **', 1)
            number = parts[0]
            heading_text = parts[1].replace('**', '')
            
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            run = paragraph.add_run(f"{number}. ")
            set_font_meiryo(run, 9)
            
            run = paragraph.add_run(heading_text)
            set_font_meiryo(run, 9)
            run.font.bold = True
            continue
        
        if not line_stripped.startswith('#'):
            paragraph = doc.add_paragraph(line_stripped)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_font_meiryo(run, 9)
    
    doc.add_page_break()

def add_colophon_page(doc, md_path):
    """奥付ページを追加（野沢温泉形式）"""
    print("  ✓ 奥付（野沢温泉形式）を追加")
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    skip_first_heading = True
    
    for line in lines:
        line_stripped = line.strip()
        
        if not line_stripped:
            doc.add_paragraph()
            continue
        
        if skip_first_heading and line_stripped.startswith('# '):
            skip_first_heading = False
            continue
        
        if line_stripped.startswith('## 著者プロフィール'):
            heading_text = line_stripped.replace('## ', '')
            paragraph = doc.add_paragraph(heading_text)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_font_meiryo(run, 11)
                run.font.bold = True
            continue
        
        if line_stripped.startswith('## ChatGPT'):
            doc.add_paragraph()
            heading_text = line_stripped.replace('## ', '')
            paragraph = doc.add_paragraph(heading_text)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_font_meiryo(run, 11)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 176, 240)
            continue
        
        if line_stripped.startswith('---'):
            continue
        
        if not line_stripped.startswith('#'):
            paragraph = doc.add_paragraph(line_stripped)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_font_meiryo(run, 9)

def main():
    print("Word原稿v30を作成中（各要素をマージ）...")
    print("=" * 80)
    
    # 新しいドキュメントを作成
    doc = Document()
    
    # 1. 扉（画像）
    print("\n1. 扉を追加中...")
    add_title_page(doc, '/home/ubuntu/ai_travel_book_project/finals/title_page_final.png')
    
    # 2. 前書き
    print("\n2. 前書きを追加中...")
    add_markdown_section(doc, '/home/ubuntu/ai_travel_book_project/finals/preface_v1.md', '前書き')
    doc.add_page_break()
    
    # 3. 目次（v27から抽出）
    print("\n3. 目次を追加中...")
    v27 = Document('/home/ubuntu/ai_travel_book_project/finals/complete_manuscript_v27_safe.docx')
    
    # v27の構造: 段落1〜154が目次、段落155から本文
    toc_start = 0
    toc_end = 154
    
    print(f"  ✓ 目次を抽出（段落{toc_start+1}〜{toc_end+1}）")
    for i in range(toc_start, toc_end + 1):
        copy_paragraph(doc, v27.paragraphs[i])
    doc.add_page_break()
    
    # 4. 中扉
    print("\n4. 中扉を追加中...")
    add_markdown_section(doc, '/home/ubuntu/ai_travel_book_project/finals/halftitle_v1.md', '中扉')
    doc.add_page_break()
    
    # 5. 本文（v27から第1章〜第12章+付録を抽出）
    print("\n5. 本文を追加中...")
    
    # v27の構造: 段落155（第1章）から免責事項の前まで
    chapter1_start = 154  # 段落155（インデックス154）
    
    # 免責事項の開始を探す（本文の終わり）
    disclaimer_start = None
    for i, para in enumerate(v27.paragraphs):
        text = para.text.strip()
        if text == '免責事項':
            disclaimer_start = i
            break
    
    if disclaimer_start:
        print(f"  ✓ 本文を抽出（段落{chapter1_start+1}〜{disclaimer_start}）")
        for i in range(chapter1_start, disclaimer_start):
            copy_paragraph(doc, v27.paragraphs[i])
        doc.add_page_break()
    else:
        print("  警告: 免責事項が見つかりませんでした、最後まで本文をコピーします")
        print(f"  ✓ 本文を抽出（段落{chapter1_start+1}〜最後）")
        for i in range(chapter1_start, len(v27.paragraphs)):
            copy_paragraph(doc, v27.paragraphs[i])
        doc.add_page_break()
    
    # 6. 後書き
    print("\n6. 後書きを追加中...")
    add_markdown_section(doc, '/home/ubuntu/ai_travel_book_project/finals/afterword_v1.md', '後書き')
    doc.add_page_break()
    
    # 7. 免責事項
    print("\n7. 免責事項を追加中...")
    add_disclaimer_page(doc, '/home/ubuntu/ai_travel_book_project/finals/disclaimer_v2.md')
    
    # 8. 奥付
    print("\n8. 奥付を追加中...")
    add_colophon_page(doc, '/home/ubuntu/ai_travel_book_project/finals/colophon_v2.md')
    
    # 保存
    output_path = '/home/ubuntu/ai_travel_book_project/finals/complete_manuscript_v30.docx'
    doc.save(output_path)
    
    print("\n" + "=" * 80)
    print(f"✓ v30原稿を保存しました: {output_path}")
    print(f"✓ 総段落数: {len(doc.paragraphs)}")
    print("=" * 80)
    print("完了！")

if __name__ == '__main__':
    main()
