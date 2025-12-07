#!/usr/bin/env python3
"""
v16の見出し構造を確認するスクリプト
"""

from docx import Document

def check_document_structure(docx_path):
    """Word文書の見出し構造を確認"""
    
    doc = Document(docx_path)
    
    print(f"文書: {docx_path}")
    print(f"段落数: {len(doc.paragraphs)}")
    print(f"\n見出し一覧:")
    print("-" * 80)
    
    heading1_count = 0
    heading2_count = 0
    heading3_count = 0
    
    for i, para in enumerate(doc.paragraphs):
        if para.style.name.startswith('Heading'):
            level = para.style.name.replace('Heading ', '')
            text = para.text[:80] if len(para.text) > 80 else para.text
            print(f"{para.style.name:15} | {text}")
            
            if para.style.name == 'Heading 1':
                heading1_count += 1
            elif para.style.name == 'Heading 2':
                heading2_count += 1
            elif para.style.name == 'Heading 3':
                heading3_count += 1
    
    print("-" * 80)
    print(f"\n見出し統計:")
    print(f"  Heading 1: {heading1_count}")
    print(f"  Heading 2: {heading2_count}")
    print(f"  Heading 3: {heading3_count}")

if __name__ == "__main__":
    check_document_structure("complete_manuscript_v16_merged.docx")
