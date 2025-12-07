#!/usr/bin/env python3
"""
v24のWord文書を結合するスクリプト
"""

from docx import Document

def merge_docx(files, output):
    """複数のWord文書を結合"""
    merged_doc = Document(files[0])
    
    for file in files[1:]:
        doc = Document(file)
        
        # ページ区切りなしで直接要素をコピー
        for element in doc.element.body:
            merged_doc.element.body.append(element)
    
    merged_doc.save(output)
    print(f"結合完了: {output}")

def count_headings(docx_file):
    """見出しの数をカウント"""
    doc = Document(docx_file)
    h1_count = 0
    h2_count = 0
    h3_count = 0
    para_count = 0
    table_count = len(doc.tables)
    
    for para in doc.paragraphs:
        para_count += 1
        if para.style.name == 'Heading 1':
            h1_count += 1
            print(f"  見出し1 ({h1_count}): {para.text[:50]}")
        elif para.style.name == 'Heading 2':
            h2_count += 1
        elif para.style.name == 'Heading 3':
            h3_count += 1
    
    return h1_count, h2_count, h3_count, para_count, table_count

# 結合するファイルのリスト（順序が重要）
files = [
    'v24_quickstart.docx',
    'v24_introduction_final.docx',
    'v24_chapter1_final.docx',
    'v24_chapter2_final.docx',
    'v24_chapter3_final.docx',
    'v24_chapter4_final.docx',
    'v24_chapter5_final.docx',
    'v24_chapter6_final.docx',
    'v24_chapter7_final.docx',
    'v24_chapter8_final.docx',
    'v24_chapter9_final.docx',
    'v24_chapter10_final.docx',
    'v24_chapter11_final.docx',
    'v24_chapter12_final.docx',
    'v24_afterword.docx',
    'v24_appendix_b_final.docx',
    'v24_appendix_c_final.docx',
]

output_file = 'complete_manuscript_v24_final.docx'

print("=== v24のWord文書を結合中 ===")
merge_docx(files, output_file)

print("\n=== 見出し構造の検証 ===")
h1, h2, h3, paras, tables = count_headings(output_file)

print(f"\n【検証結果】")
print(f"見出し1の数: {h1} (期待値: 17)")
print(f"見出し2の数: {h2}")
print(f"見出し3の数: {h3}")
print(f"総段落数: {paras:,}")
print(f"テーブル数: {tables}")

if h1 == 17:
    print("\n✅ 見出し1の数が正しいです！")
else:
    print(f"\n⚠️ 見出し1の数が期待値と異なります（期待: 17, 実際: {h1}）")
