#!/usr/bin/env python3
"""
v16_finalの見出し構造を確認するスクリプト
"""

from docx import Document

def check_document_structure(docx_path):
    """Word文書の見出し構造を確認"""
    
    doc = Document(docx_path)
    
    print(f"文書: {docx_path}")
    print(f"段落数: {len(doc.paragraphs)}")
    print(f"\nHeading 1の一覧:")
    print("-" * 80)
    
    heading1_list = []
    heading1_count = 0
    heading2_count = 0
    heading3_count = 0
    
    for i, para in enumerate(doc.paragraphs):
        if para.style.name == 'Heading 1':
            text = para.text[:80] if len(para.text) > 80 else para.text
            print(f"{heading1_count + 1:2}. {text}")
            heading1_list.append(text)
            heading1_count += 1
        elif para.style.name == 'Heading 2':
            heading2_count += 1
        elif para.style.name == 'Heading 3':
            heading3_count += 1
    
    print("-" * 80)
    print(f"\n見出し統計:")
    print(f"  Heading 1: {heading1_count}")
    print(f"  Heading 2: {heading2_count}")
    print(f"  Heading 3: {heading3_count}")
    
    # 期待される17個のHeading 1
    expected_h1 = [
        "クイックスタートガイド",
        "はじめに",
        "第1章",
        "第2章",
        "第3章",
        "第4章",
        "第5章",
        "第6章",
        "第7章",
        "第8章",
        "第9章",
        "第10章",
        "第11章",
        "第12章",
        "おわりに",
        "付録B",
        "付録C"
    ]
    
    print(f"\n期待されるHeading 1: {len(expected_h1)}")
    print(f"実際のHeading 1: {heading1_count}")
    
    if heading1_count == len(expected_h1):
        print("\n✓ すべてのHeading 1が存在します")
    else:
        print(f"\n✗ Heading 1が {len(expected_h1) - heading1_count} 個不足しています")

if __name__ == "__main__":
    check_document_structure("complete_manuscript_v16_final.docx")
