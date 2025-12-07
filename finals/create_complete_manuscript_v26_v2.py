#!/usr/bin/env python3
"""
完全なWord原稿（v26）を作成する自動化スクリプト（v2）

新しいアプローチ:
1. 各MDファイルを個別にWord文書に変換
2. python-docxで全Word文書を結合
3. 欠落テーブルの位置にプレースホルダーを挿入
4. python-docxで各プレースホルダーにテーブルを挿入
"""

import os
import re
import json
import subprocess
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# === 設定 ===
WORK_DIR = '/home/ubuntu/ai_travel_book_project/finals'
os.chdir(WORK_DIR)

# MDファイルのリスト（順序重要）
MD_FILES = [
    'v21_part_00_quickstart.md',
    'introduction_final.md',
    'chapter1_final.md',
    'chapter2_final.md',
    'chapter3_final.md',
    'chapter4_final.md',
    'chapter5_final.md',
    'chapter6_final.md',
    'chapter7_final.md',
    'chapter8_final.md',
    'chapter9_final.md',
    'chapter10_final.md',
    'chapter11_final.md',
    'chapter12_final.md',
    'conclusion_final.md',
    'appendix_b_final.md',
    'appendix_c_final.md'
]

# 欠落テーブルの情報
MISSING_TABLES = [
    {'file': 'chapter3_final.md', 'line': 69, 'id': 'ch3_t1'},
    {'file': 'chapter3_final.md', 'line': 140, 'id': 'ch3_t2'},
    {'file': 'chapter3_final.md', 'line': 147, 'id': 'ch3_t3'},
    {'file': 'chapter3_final.md', 'line': 156, 'id': 'ch3_t4'},
    {'file': 'chapter3_final.md', 'line': 283, 'id': 'ch3_t5'},
    {'file': 'chapter3_final.md', 'line': 343, 'id': 'ch3_t6'},
    {'file': 'chapter4_final.md', 'line': 214, 'id': 'ch4_t1'},
    {'file': 'chapter4_final.md', 'line': 221, 'id': 'ch4_t2'},
    {'file': 'chapter4_final.md', 'line': 249, 'id': 'ch4_t3'},
    {'file': 'chapter4_final.md', 'line': 307, 'id': 'ch4_t4'}
]

FINAL_DOCX = 'complete_manuscript_v26_final.docx'
TABLE_DATA_JSON = 'missing_tables_data_cleaned.json'

print("=" * 60)
print("完全なWord原稿（v26）作成スクリプト v2")
print("=" * 60)

# === ステップ1: 各MDファイルにプレースホルダーを挿入 ===
print("\n【ステップ1】各MDファイルにプレースホルダーを挿入中...")

file_contents = {}
for md_file in MD_FILES:
    if os.path.exists(md_file):
        with open(md_file, 'r', encoding='utf-8') as f:
            file_contents[md_file] = f.read()

# 欠落テーブルをプレースホルダーに置き換え
for table_info in MISSING_TABLES:
    filename = table_info['file']
    target_line = table_info['line']
    table_id = table_info['id']
    
    if filename not in file_contents:
        continue
    
    lines = file_contents[filename].split('\n')
    
    # テーブルの開始行と終了行を特定
    table_start = target_line - 1  # 0-indexed
    table_end = table_start
    
    for i in range(table_start, len(lines)):
        if re.match(r'^\|.+\|$', lines[i]):
            table_end = i
        else:
            break
    
    # テーブルをプレースホルダーに置き換え
    placeholder = f'【テーブル挿入位置:{table_id}】'
    lines[table_start:table_end+1] = [placeholder]
    
    file_contents[filename] = '\n'.join(lines)
    print(f"  ✅ {table_id}: {filename} 行{target_line}")

# プレースホルダー付きのMDファイルを保存
temp_dir = 'temp_md_files'
os.makedirs(temp_dir, exist_ok=True)

for md_file in MD_FILES:
    if md_file in file_contents:
        temp_file = os.path.join(temp_dir, md_file)
        with open(temp_file, 'w', encoding='utf-8') as f:
            f.write(file_contents[md_file])

print(f"✅ プレースホルダー挿入完了")

# === ステップ2: 各MDファイルを個別にWord文書に変換 ===
print("\n【ステップ2】各MDファイルをWord文書に変換中...")

temp_docx_dir = 'temp_docx_files'
os.makedirs(temp_docx_dir, exist_ok=True)

docx_files = []
for md_file in MD_FILES:
    temp_md = os.path.join(temp_dir, md_file)
    if not os.path.exists(temp_md):
        print(f"  ⚠️  {md_file} が見つかりません。スキップします。")
        continue
    
    docx_file = os.path.join(temp_docx_dir, md_file.replace('.md', '.docx'))
    
    pandoc_cmd = [
        'pandoc',
        temp_md,
        '-o', docx_file,
        '--from=markdown',
        '--to=docx'
    ]
    
    result = subprocess.run(pandoc_cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"  ❌ {md_file} の変換エラー: {result.stderr}")
        continue
    
    docx_files.append(docx_file)
    print(f"  ✅ {md_file} → {os.path.basename(docx_file)}")

print(f"✅ {len(docx_files)}個のWord文書を作成しました")

# === ステップ3: 全Word文書を結合 ===
print("\n【ステップ3】全Word文書を結合中...")

# 最初の文書を基準にする
combined_doc = Document(docx_files[0])

# 2番目以降の文書を追加
for docx_file in docx_files[1:]:
    sub_doc = Document(docx_file)
    
    # ページブレークを追加
    combined_doc.add_page_break()
    
    # 段落とテーブルを追加
    for element in sub_doc.element.body:
        combined_doc.element.body.append(element)
    
    print(f"  ✅ {os.path.basename(docx_file)} を結合")

# 一旦保存
temp_combined = 'temp_combined.docx'
combined_doc.save(temp_combined)
print(f"✅ 結合完了: {temp_combined}")

# === ステップ4: テーブルを手動で挿入 ===
print("\n【ステップ4】テーブルを手動で挿入中...")

# テーブルデータを読み込み
with open(TABLE_DATA_JSON, 'r', encoding='utf-8') as f:
    table_data_collection = json.load(f)

# Word文書を開く
doc = Document(temp_combined)

# 各プレースホルダーを検索してテーブルを挿入
for table_info in MISSING_TABLES:
    table_id = table_info['id']
    placeholder = f'【テーブル挿入位置:{table_id}】'
    
    if table_id not in table_data_collection:
        print(f"  ⚠️  警告: {table_id} のデータが見つかりません。")
        continue
    
    table_data = table_data_collection[table_id]['data']
    
    # プレースホルダーを検索
    found = False
    for para in doc.paragraphs:
        if placeholder in para.text:
            # テーブルを作成
            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            table.autofit = True
            
            # テーブルにデータを入力
            for i, row_data in enumerate(table_data):
                for j, cell_data in enumerate(row_data):
                    cell = table.rows[i].cells[j]
                    cell.text = cell_data
                    # フォントサイズを設定
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
            
            # プレースホルダーの段落をテーブルに置き換え
            parent = para._p.getparent()
            parent.insert(parent.index(para._p), table._element)
            parent.remove(para._p)
            
            print(f"  ✅ {table_id}: {len(table_data)}行×{len(table_data[0])}列のテーブルを挿入")
            found = True
            break
    
    if not found:
        print(f"  ⚠️  警告: {table_id} のプレースホルダーが見つかりません。")

# 文書を保存
doc.save(FINAL_DOCX)
print(f"\n✅ テーブル挿入完了: {FINAL_DOCX}")

# === ステップ5: 統計情報を表示 ===
print("\n【ステップ5】統計情報")

doc_final = Document(FINAL_DOCX)
print(f"  総段落数: {len(doc_final.paragraphs)}")
print(f"  総テーブル数: {len(doc_final.tables)}")

# Heading 1の数を数える
h1_count = sum(1 for p in doc_final.paragraphs if p.style.name == 'Heading 1')
print(f"  Heading 1の数: {h1_count}")

print("\n【Heading 1のリスト】")
for p in doc_final.paragraphs:
    if p.style.name == 'Heading 1':
        print(f"  - {p.text}")

print("\n" + "=" * 60)
print("✅ 完全なWord原稿（v26）の作成が完了しました！")
print(f"出力ファイル: {FINAL_DOCX}")
print("=" * 60)
