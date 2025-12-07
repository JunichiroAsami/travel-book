#!/usr/bin/env python3
"""
Word文書から不要なページブレークを削除するスクリプト
"""

from docx import Document
from docx.oxml.ns import qn

INPUT_DOCX = 'complete_manuscript_v26_with_headers.docx'
OUTPUT_DOCX = 'complete_manuscript_v26_no_breaks.docx'

print("=" * 60)
print("ページブレーク削除スクリプト")
print("=" * 60)

# Word文書を開く
doc = Document(INPUT_DOCX)

print(f"\n総段落数: {len(doc.paragraphs)}")

# ページブレークを削除
removed_count = 0

for para in doc.paragraphs:
    for run in para.runs:
        # ページブレークを検出
        br_elements = run._element.findall(qn('w:br'))
        
        for br in br_elements:
            # ページブレークの種類を確認
            br_type = br.get(qn('w:type'))
            
            # ページブレークを削除
            if br_type == 'page':
                br.getparent().remove(br)
                removed_count += 1

print(f"\n削除したページブレーク数: {removed_count}")

# 保存
doc.save(OUTPUT_DOCX)

print(f"\n✅ 完了: {OUTPUT_DOCX}")
print("=" * 60)
