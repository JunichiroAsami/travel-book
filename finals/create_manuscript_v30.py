#!/usr/bin/env python3
"""Word原稿v30を作成: 野沢温泉形式の免責事項と奥付を適用"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font_meiryo(run, size_pt):
    """フォントをメイリオに設定"""
    run.font.name = 'メイリオ'
    run.font.size = Pt(size_pt)
    # 英数字もメイリオにする
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), 'メイリオ')
    rFonts.set(qn('w:ascii'), 'メイリオ')
    rFonts.set(qn('w:hAnsi'), 'メイリオ')

def add_disclaimer_page(doc, md_path):
    """免責事項ページを追加（野沢温泉形式）"""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    skip_first_heading = True
    in_numbered_section = False
    
    for line in lines:
        line_stripped = line.strip()
        
        # 空行
        if not line_stripped:
            doc.add_paragraph()
            continue
        
        # 最初の見出し（# 免責事項）はスキップ
        if skip_first_heading and line_stripped.startswith('# '):
            skip_first_heading = False
            
            # 見出しを追加
            heading = doc.add_heading('免責事項（ディスクレーマー）', level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in heading.runs:
                set_font_meiryo(run, 12)
                run.font.bold = True
            continue
        
        # 番号付きリスト（1. 2. 3. ...）の検出
        if line_stripped[0].isdigit() and '. **' in line_stripped:
            in_numbered_section = True
            
            # 番号と見出しを分離
            parts = line_stripped.split('. **', 1)
            number = parts[0]
            heading_text = parts[1].replace('**', '')
            
            # 段落を追加
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 番号を追加
            run = paragraph.add_run(f"{number}. ")
            set_font_meiryo(run, 9)
            
            # 見出しを追加（太字）
            run = paragraph.add_run(heading_text)
            set_font_meiryo(run, 9)
            run.font.bold = True
            
            continue
        
        # 通常のテキスト
        if not line_stripped.startswith('#'):
            paragraph = doc.add_paragraph(line_stripped)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_font_meiryo(run, 9)
    
    doc.add_page_break()

def add_colophon_page(doc, md_path):
    """奥付ページを追加（野沢温泉形式）"""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    skip_first_heading = True
    in_author_profile = False
    in_colophon = False
    
    for line in lines:
        line_stripped = line.strip()
        
        # 空行
        if not line_stripped:
            doc.add_paragraph()
            continue
        
        # 最初の見出し（# 奥付）はスキップ
        if skip_first_heading and line_stripped.startswith('# '):
            skip_first_heading = False
            continue
        
        # 著者プロフィールの見出し
        if line_stripped.startswith('## 著者プロフィール'):
            in_author_profile = True
            
            # 見出しを追加
            heading_text = line_stripped.replace('## ', '')
            paragraph = doc.add_paragraph(heading_text)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_font_meiryo(run, 11)
                run.font.bold = True
            continue
        
        # 書籍タイトルの見出し（青色）
        if line_stripped.startswith('## ChatGPT'):
            in_author_profile = False
            in_colophon = True
            
            # 区切り線（---）の後なので空行を追加
            doc.add_paragraph()
            
            # 見出しを追加（青色）
            heading_text = line_stripped.replace('## ', '')
            paragraph = doc.add_paragraph(heading_text)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_font_meiryo(run, 11)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 176, 240)  # 青色
            continue
        
        # 区切り線（---）
        if line_stripped.startswith('---'):
            continue
        
        # 通常のテキスト
        if not line_stripped.startswith('#'):
            paragraph = doc.add_paragraph(line_stripped)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # フォントサイズを設定
            if in_author_profile:
                for run in paragraph.runs:
                    set_font_meiryo(run, 9)
            elif in_colophon:
                for run in paragraph.runs:
                    set_font_meiryo(run, 9)

def main():
    print("Word原稿v30を作成中（野沢温泉形式の免責事項と奥付）...")
    print("=" * 80)
    
    # v29を読み込む
    print("✓ v29原稿を読み込み中...")
    doc = Document('/home/ubuntu/ai_travel_book_project/finals/complete_manuscript_v29.docx')
    
    # v29の最後の2ページ（古い免責事項と奥付）を削除
    # 注: python-docxでは直接ページを削除できないため、
    # 最後の数段落を削除する方法で対応
    print("✓ 古い免責事項と奥付を削除中...")
    
    # 最後の段落から「免責事項」または「奥付」を含む段落を探して削除
    paragraphs_to_remove = []
    found_disclaimer = False
    
    for i in range(len(doc.paragraphs) - 1, -1, -1):
        paragraph = doc.paragraphs[i]
        text = paragraph.text.strip()
        
        # 「免責事項」または「奥付」が見つかったら、そこから最後まで削除
        if '免責事項' in text or '奥付' in text or found_disclaimer:
            found_disclaimer = True
            paragraphs_to_remove.append(paragraph)
    
    # 削除
    for paragraph in paragraphs_to_remove:
        p = paragraph._element
        p.getparent().remove(p)
    
    # 新しい免責事項を追加
    print("✓ 野沢温泉形式の免責事項を追加中...")
    add_disclaimer_page(doc, '/home/ubuntu/ai_travel_book_project/finals/disclaimer_v2.md')
    
    # 新しい奥付を追加
    print("✓ 野沢温泉形式の奥付を追加中...")
    add_colophon_page(doc, '/home/ubuntu/ai_travel_book_project/finals/colophon_v2.md')
    
    # 保存
    output_path = '/home/ubuntu/ai_travel_book_project/finals/complete_manuscript_v30.docx'
    doc.save(output_path)
    print(f"✓ v30原稿を保存しました: {output_path}")
    
    print("=" * 80)
    print("完了！")
    print("\n注意: フォントと行間の微調整はWordで手動で行ってください。")

if __name__ == '__main__':
    main()
