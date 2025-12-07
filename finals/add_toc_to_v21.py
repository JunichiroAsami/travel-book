#!/usr/bin/env python3.11
"""
v21のWord文書に目次を追加
"""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

print("=" * 80)
print("v21に目次を追加")
print("=" * 80)
print()

# v21を読み込み
doc = Document("complete_manuscript_v21_final.docx")
print(f"✓ v21を読み込みました（{len(doc.paragraphs)}段落）")

# 目次を最初に挿入
# 新しい段落を最初に挿入するため、逆順で追加

# 1. ページ区切りを追加（目次の後）
page_break = doc.paragraphs[0]._element
page_break_elem = OxmlElement('w:br')
page_break_elem.set(qn('w:type'), 'page')
page_break._element.insert(0, page_break_elem)

# 2. 目次フィールドを追加
toc_para = doc.paragraphs[0]._element
toc_field = OxmlElement('w:sdt')

# 目次のXMLを構築
sdtPr = OxmlElement('w:sdtPr')
sdtContent = OxmlElement('w:sdtContent')

# 目次フィールドコード
fldSimple = OxmlElement('w:fldSimple')
fldSimple.set(qn('w:instr'), 'TOC \\o "1-2" \\h \\z \\u')

sdtContent.append(fldSimple)
toc_field.append(sdtPr)
toc_field.append(sdtContent)

# 最初の段落の前に挿入
doc._element.body.insert(0, toc_field)

print("✓ 目次フィールドを追加しました")

# 3. 目次タイトルを追加
toc_title = doc.paragraphs[0]._element
toc_title_para = OxmlElement('w:p')

# タイトルのテキスト
toc_title_run = OxmlElement('w:r')
toc_title_text = OxmlElement('w:t')
toc_title_text.text = "TABLE OF CONTENTS"
toc_title_run.append(toc_title_text)

# フォントサイズとボールド
rPr = OxmlElement('w:rPr')
sz = OxmlElement('w:sz')
sz.set(qn('w:val'), '32')  # 16pt = 32 half-points
szCs = OxmlElement('w:szCs')
szCs.set(qn('w:val'), '32')
b = OxmlElement('w:b')
rPr.append(sz)
rPr.append(szCs)
rPr.append(b)
toc_title_run.insert(0, rPr)

toc_title_para.append(toc_title_run)

# 中央揃え
pPr = OxmlElement('w:pPr')
jc = OxmlElement('w:jc')
jc.set(qn('w:val'), 'center')
pPr.append(jc)

# 背景色を青に設定
shd = OxmlElement('w:shd')
shd.set(qn('w:fill'), '00B0F0')
pPr.append(shd)

toc_title_para.insert(0, pPr)

# 最初に挿入
doc._element.body.insert(0, toc_title_para)

print("✓ 目次タイトルを追加しました")

# 保存
output_file = "complete_manuscript_v21_final.docx"
doc.save(output_file)

print()
print("=" * 80)
print(f"✓ 完了: {output_file}")
print("=" * 80)
print()

print("次のステップ:")
print("  1. Wordで開く")
print("  2. 目次を右クリック → 「フィールドの更新」→「目次をすべて更新」")
print("  3. 保存")
