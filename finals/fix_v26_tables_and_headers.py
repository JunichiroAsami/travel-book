#!/usr/bin/env python3
"""
v26原稿を修正するスクリプト:
1. テーブルの余計な文字（↩など）を削除
2. ヘッダー・フッターをv24と同じ形式に変更
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

def add_page_number(paragraph):
    """段落にページ番号を追加"""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    
    return run

def clean_table_cell_text(text):
    """テーブルセルから余計な文字を削除"""
    # ↩ (U+21A9) や他の制御文字を削除
    cleaned = text.replace('↩', '').replace('\r', '').replace('\n', ' ')
    # 複数のスペースを1つに
    cleaned = re.sub(r'\s+', ' ', cleaned)
    return cleaned.strip()

# === 設定 ===
INPUT_DOCX = 'complete_manuscript_v26_with_headers.docx'
OUTPUT_DOCX = 'complete_manuscript_v26_fixed.docx'

print("=" * 60)
print("v26原稿修正スクリプト")
print("=" * 60)

# Word文書を開く
doc = Document(INPUT_DOCX)

# === ステップ1: テーブルの余計な文字を削除 ===
print("\n【ステップ1】テーブルの余計な文字を削除中...")

table_count = 0
cell_count = 0

for table in doc.tables:
    table_count += 1
    for row in table.rows:
        for cell in row.cells:
            original_text = cell.text
            cleaned_text = clean_table_cell_text(original_text)
            
            if original_text != cleaned_text:
                # セル内の全段落をクリア
                for para in cell.paragraphs:
                    para.clear()
                
                # 最初の段落に清潔なテキストを設定
                if cell.paragraphs:
                    cell.paragraphs[0].text = cleaned_text
                else:
                    cell.add_paragraph(cleaned_text)
                
                cell_count += 1

print(f"  ✅ {table_count}個のテーブルを処理")
print(f"  ✅ {cell_count}個のセルをクリーニング")

# === ステップ2: ヘッダー・フッターを修正 ===
print("\n【ステップ2】ヘッダー・フッターを修正中...")

# 最初のセクションのヘッダーを取得
section = doc.sections[0]
header = section.header

# ヘッダーの既存の内容をクリア
for paragraph in header.paragraphs:
    paragraph.clear()

# ヘッダーにテーブルを作成（v24スタイル）
# 左側: 章タイトル、右側: ページ番号（青い背景）
from docx.shared import Inches
header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
header_table.autofit = True

# 左セル: 章タイトル（動的に変更されるため、プレースホルダー）
left_cell = header_table.rows[0].cells[0]
left_para = left_cell.paragraphs[0]
left_para.text = "野沢温泉探訪記～高市政権の外国人政策の行く末・雪国で考える「共生の設計」"
left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
left_para.runs[0].font.size = Pt(9)

# 右セル: ページ番号（青い背景）
right_cell = header_table.rows[0].cells[1]
right_para = right_cell.paragraphs[0]
right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# ページ番号を追加
add_page_number(right_para)
right_para.runs[0].font.size = Pt(11)
right_para.runs[0].font.color.rgb = RGBColor(255, 255, 255)  # 白色

# セルの背景色を青に設定
shading_elm = OxmlElement('w:shd')
shading_elm.set(qn('w:fill'), '7CB8C5')  # 青色
right_cell._element.get_or_add_tcPr().append(shading_elm)

print("  ✅ ヘッダーをv24スタイルに変更")

# フッターをクリア
footer = section.footer
for paragraph in footer.paragraphs:
    paragraph.clear()

print("  ✅ フッターをクリア")

# === ステップ3: 表紙ページの設定 ===
print("\n【ステップ3】表紙ページの設定...")

section.different_first_page_header_footer = True

# 最初のページのヘッダーとフッターをクリア
first_page_header = section.first_page_header
first_page_footer = section.first_page_footer

for paragraph in first_page_header.paragraphs:
    paragraph.clear()

for paragraph in first_page_footer.paragraphs:
    paragraph.clear()

print("  ✅ 表紙ページのヘッダー・フッターをクリア")

# === 文書を保存 ===
doc.save(OUTPUT_DOCX)

print("\n" + "=" * 60)
print(f"✅ v26原稿の修正が完了しました！")
print(f"出力ファイル: {OUTPUT_DOCX}")
print("=" * 60)
