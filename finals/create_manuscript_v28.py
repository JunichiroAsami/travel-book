#!/usr/bin/env python3
"""Word原稿v28を作成: タイトル変更、扉画像、ディスクレーマー、奥付を統合"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_break(doc):
    """改ページを追加"""
    doc.add_page_break()

def add_title_page_image(doc, image_path):
    """扉画像を追加"""
    # 新しいセクションとして扉を追加
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 画像を追加（A4サイズに合わせて調整）
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(6.5))
    
    add_page_break(doc)

def add_disclaimer(doc, disclaimer_md_path):
    """ディスクレーマーを追加"""
    with open(disclaimer_md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # タイトルを追加
    heading = doc.add_heading('免責事項', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 内容を解析して追加
    lines = content.split('\n')
    skip_first_heading = True
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # 最初の見出し（# 免責事項...）はスキップ
        if skip_first_heading and line.startswith('# '):
            skip_first_heading = False
            continue
        
        # 見出し2（## 1. ...）
        if line.startswith('## '):
            heading_text = line[3:]
            heading = doc.add_heading(heading_text, level=2)
        # 通常のテキスト
        elif not line.startswith('#'):
            paragraph = doc.add_paragraph(line)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    add_page_break(doc)

def add_colophon(doc, colophon_md_path):
    """奥付を追加"""
    with open(colophon_md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # タイトルを追加
    heading = doc.add_heading('奥付', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 内容を解析して追加
    lines = content.split('\n')
    skip_first_heading = True
    in_book_title = False
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # 最初の見出し（# 奥付...）はスキップ
        if skip_first_heading and line.startswith('# '):
            skip_first_heading = False
            continue
        
        # 見出し2（## ChatGPTと旅する...）
        if line.startswith('## '):
            book_title = line[3:]
            paragraph = doc.add_paragraph(book_title)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.runs[0]
            run.font.size = Pt(14)
            run.font.bold = True
            in_book_title = True
            doc.add_paragraph()  # 空行
        # 太字項目（**著者名**: ...）
        elif '**' in line:
            paragraph = doc.add_paragraph()
            # **で囲まれた部分を太字に
            parts = line.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 1:  # 奇数番目は太字
                    run = paragraph.add_run(part)
                    run.font.bold = True
                else:
                    run = paragraph.add_run(part)
        # 区切り線（---）
        elif line.startswith('---'):
            doc.add_paragraph()  # 空行
        # 通常のテキスト
        elif not line.startswith('#'):
            paragraph = doc.add_paragraph(line)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

def update_title_in_document(doc, new_title, new_subtitle):
    """ドキュメント内のタイトルを更新"""
    # 最初のページの見出しを探して更新
    for paragraph in doc.paragraphs[:10]:  # 最初の10段落を確認
        if paragraph.style.name.startswith('Heading') or paragraph.style.name == 'Title':
            # 古いタイトルを新しいタイトルに置き換え
            if 'AIツアーコンダクター' in paragraph.text or 'ベトナム・タイ周遊記' in paragraph.text:
                paragraph.text = new_title
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(18)
                    run.font.bold = True
                
                # サブタイトルを追加
                subtitle_para = paragraph.insert_paragraph_before(new_subtitle)
                subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in subtitle_para.runs:
                    run.font.size = Pt(14)
                break

def main():
    print("Word原稿v28を作成中...")
    print("=" * 80)
    
    # v27原稿を読み込む
    doc = Document('/home/ubuntu/ai_travel_book_project/finals/complete_manuscript_v27_safe.docx')
    print("✓ v27原稿を読み込みました")
    
    # 新しいドキュメントを作成
    new_doc = Document()
    
    # 1. 扉画像を追加
    print("✓ 扉画像を追加中...")
    add_title_page_image(new_doc, '/home/ubuntu/ai_travel_book_project/finals/title_page_final.png')
    
    # 2. v27の内容をコピー（タイトルは更新）
    print("✓ v27の内容をコピー中...")
    for element in doc.element.body:
        new_doc.element.body.append(element)
    
    # 3. タイトルを更新
    print("✓ タイトルを更新中...")
    new_title = "ChatGPTと旅するベトナム・タイ"
    new_subtitle = "AI活用で言葉の壁を越える実践ガイド"
    update_title_in_document(new_doc, new_title, new_subtitle)
    
    # 4. ディスクレーマーを追加
    print("✓ ディスクレーマーを追加中...")
    add_disclaimer(new_doc, '/home/ubuntu/ai_travel_book_project/finals/disclaimer_v1.md')
    
    # 5. 奥付を追加
    print("✓ 奥付を追加中...")
    add_colophon(new_doc, '/home/ubuntu/ai_travel_book_project/finals/colophon_v1.md')
    
    # 6. 保存
    output_path = '/home/ubuntu/ai_travel_book_project/finals/complete_manuscript_v28.docx'
    new_doc.save(output_path)
    print(f"✓ v28原稿を保存しました: {output_path}")
    
    print("=" * 80)
    print("完了！")

if __name__ == '__main__':
    main()
