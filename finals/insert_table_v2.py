from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# chapter3のWord文書を読み込み
doc = Document('chapter3_no_table_v6.docx')

# テーブルデータ
table_data = [
    ['ホテル名', '予算', '朝食', 'プール', 'ルーフトップバー', 'トゥクトゥク乗車'],
    ['Avani+ Riverside', '◎ 約1万円台', '◎ 豊富＋景観', '◎ インフィニティプール', '◎ プール併設バー', '○ 川沿い移動可'],
    ['Eastin Grand Sathorn', '◎ リーズナブル', '◎ フルーツ・オムレツ', '◎ 屋外プール', '○ 屋外バーあり', '○ 近隣で利用可'],
    ['Marriott Surawongse', '○ プロモ価格で可', '◎ 高評価ビュッフェ', '◎ インフィニティプール', '◎ 33階ルーフトップバー', '○ シーロム界隈で可'],
    ['Sivatel バンコク', '◎ コスパ高評価', '○ 種類少なめも質あり', '◎ 静かな屋外プール', '× 記載なし', '○ BTS周辺で可']
]

# プレースホルダーを探す
placeholder_index = None
for i, para in enumerate(doc.paragraphs):
    if '【テーブル挿入位置】' in para.text:
        placeholder_index = i
        print(f"✅ プレースホルダーを発見: 段落{i}")
        break

if placeholder_index is None:
    print("⚠️ プレースホルダーが見つかりませんでした")
    exit(1)

# プレースホルダーの段落を削除して、その位置にテーブルを挿入
placeholder_para = doc.paragraphs[placeholder_index]
placeholder_element = placeholder_para._element
parent = placeholder_element.getparent()
placeholder_position = list(parent).index(placeholder_element)

# プレースホルダーを削除
parent.remove(placeholder_element)

# テーブルを作成
table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))

# テーブルにデータを入力
for i, row_data in enumerate(table_data):
    row = table.rows[i]
    for j, cell_data in enumerate(row_data):
        cell = row.cells[j]
        cell.text = cell_data
        
        # ヘッダー行は太字
        if i == 0:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

# テーブルを正しい位置に移動
table_element = table._element
parent.insert(placeholder_position, table_element)

# 保存
doc.save('chapter3_final_complete.docx')

print(f"✅ テーブルを挿入しました: chapter3_final_complete.docx")

# 検証
doc2 = Document('chapter3_final_complete.docx')
print(f"検証: 総テーブル数 = {len(doc2.tables)}")
for i, table in enumerate(doc2.tables):
    print(f"  テーブル{i+1}: {len(table.rows)}行 × {len(table.rows[0].cells)}列")
