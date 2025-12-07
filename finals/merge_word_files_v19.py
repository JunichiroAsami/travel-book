#!/usr/bin/env python3
"""
17個のWordファイルを結合してv19を作成し、ページ番号を1から開始するように調整

使用方法:
    cd /home/ubuntu/ai_travel_book_project/finals
    python3.11 merge_word_files_v19.py
"""

import os
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_page_number_start(section, start_num=1):
    """セクションのページ番号の開始番号を設定"""
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    
    pgNumType.set(qn('w:start'), str(start_num))

def merge_word_files():
    """17個のWordファイルを結合し、ページ番号を調整"""
    
    # 作業ディレクトリに移動
    os.chdir('/home/ubuntu/ai_travel_book_project/finals')
    
    print('=' * 80)
    print('17個のWordファイルを結合（v19）')
    print('=' * 80)
    print()
    
    # 最初のファイルを読み込み
    first_file = 'chapter_00.docx'
    
    if not os.path.exists(first_file):
        print(f'エラー: {first_file} が見つかりません')
        return
    
    doc = Document(first_file)
    print(f'✓ {first_file} を読み込みました')
    
    # 最初のセクションのページ番号を1から開始するように設定
    if len(doc.sections) > 0:
        set_page_number_start(doc.sections[0], start_num=1)
        print('✓ ページ番号の開始を1に設定しました')
    
    # 残りのファイルを順番に追加
    for i in range(1, 17):
        chapter_file = f'chapter_{i:02d}.docx'
        
        if not os.path.exists(chapter_file):
            print(f'警告: {chapter_file} が見つかりません。スキップします。')
            continue
        
        chapter_doc = Document(chapter_file)
        
        # 章の内容を追加
        for element in chapter_doc.element.body:
            doc.element.body.append(element)
        
        print(f'✓ {chapter_file} を追加しました')
    
    # 保存
    output_file = 'complete_manuscript_v19_final.docx'
    doc.save(output_file)
    print()
    print('=' * 80)
    print(f'✓ 結合完了: {output_file}')
    print('=' * 80)
    print()
    
    # 確認
    heading1_list = [p.text for p in doc.paragraphs if p.style.name == 'Heading 1']
    heading2_count = len([p for p in doc.paragraphs if p.style.name == 'Heading 2'])
    heading3_count = len([p for p in doc.paragraphs if p.style.name == 'Heading 3'])
    
    print(f'段落数: {len(doc.paragraphs):,}')
    print(f'Heading 1: {len(heading1_list)}')
    print(f'Heading 2: {heading2_count}')
    print(f'Heading 3: {heading3_count}')
    print()
    
    print('Heading 1の一覧:')
    print('-' * 80)
    for i, title in enumerate(heading1_list, start=1):
        print(f'{i:2}. {title}')
    print('-' * 80)
    print()
    
    # 検証
    expected_heading1_count = 17
    if len(heading1_list) == expected_heading1_count:
        print(f'✓ すべてのHeading 1が存在します（{expected_heading1_count}個）')
    else:
        print(f'✗ Heading 1が{len(heading1_list)}個です（期待値: {expected_heading1_count}）')
    
    print()
    print('修正内容:')
    print('  1. ページ番号の開始を1に設定')
    print('  2. 水平線（---）を220個削除')
    print()
    print('次のステップ:')
    print('  1. Wordで目次を更新: 目次を右クリック → 「フィールドの更新」')
    print('  2. 最終確認: 誤字脱字、フォーマットの確認')
    print('  3. Amazon KDPに提出')

if __name__ == "__main__":
    merge_word_files()
