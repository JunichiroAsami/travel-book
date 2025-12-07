#!/usr/bin/env python3
"""
改善版v27作成スクリプト

レビューで指摘された以下の改善策を実装：
1. 本文レイアウト：段落の一字下げ、見出しスタイルの明確化
2. テーブル：枠線を2.0ptに、図表番号とキャプションを追加
3. フォントと行間：統一感を持たせる
"""

import os
import json
import subprocess
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# MDファイルのリスト
md_files = [
    "v21_part_00_quickstart.md",
    "introduction_final.md",
    "chapter1_final.md",
    "chapter2_final.md",
    "chapter3_final.md",
    "chapter4_final.md",
    "chapter5_final.md",
    "chapter6_final.md",
    "chapter7_final.md",
    "chapter8_final.md",
    "chapter9_final.md",
    "chapter10_final.md",
    "chapter11_final.md",
    "chapter12_final.md",
    "conclusion_final.md",
    "appendix_b_final.md",
    "appendix_c_final.md"
]

# 欠落テーブルデータを読み込み
with open('missing_tables_data_cleaned.json', 'r', encoding='utf-8') as f:
    missing_tables = json.load(f)

print("="*60)
print("改善版v27作成スクリプト")
print("="*60)

# ステップ1: 各MDファイルを個別にWord文書に変換
print("\nステップ1: 各MDファイルをWord文書に変換")
os.makedirs("temp_md_files_v27", exist_ok=True)
os.makedirs("temp_docx_files_v27", exist_ok=True)

for md_file in md_files:
    if not os.path.exists(md_file):
        print(f"  ⚠️  {md_file} が見つかりません。スキップします。")
        continue
    
    # MDファイルをコピー
    temp_md_path = f"temp_md_files_v27/{md_file}"
    subprocess.run(['cp', md_file, temp_md_path], check=True)
    
    # Pandocで変換
    temp_docx_path = f"temp_docx_files_v27/{md_file.replace('.md', '.docx')}"
    pandoc_cmd = [
        'pandoc',
        temp_md_path,
        '-o', temp_docx_path,
        '--standalone'
    ]
    subprocess.run(pandoc_cmd, check=True)
    print(f"  ✅ {md_file} → {temp_docx_path}")

# ステップ2: 全Word文書を結合
print("\nステップ2: 全Word文書を結合")
combined_doc = Document()

# 最初のセクションの設定
section = combined_doc.sections[0]
section.page_height = Inches(11.69)  # A4
section.page_width = Inches(8.27)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

for md_file in md_files:
    temp_docx_path = f"temp_docx_files_v27/{md_file.replace('.md', '.docx')}"
    if not os.path.exists(temp_docx_path):
        continue
    
    doc = Document(temp_docx_path)
    for element in doc.element.body:
        combined_doc.element.body.append(element)
    
    print(f"  ✅ {md_file} を結合")

# 一時保存
combined_doc.save('temp_combined_v27.docx')
print("\n✅ 結合完了: temp_combined_v27.docx")

# ステップ3: スタイルの改善
print("\nステップ3: スタイルの改善")
doc = Document('temp_combined_v27.docx')

# 本文スタイルの設定
for style_name in ['Normal', 'Body Text', 'First Paragraph']:
    if style_name in doc.styles:
        style = doc.styles[style_name]
        style.font.name = 'MS Mincho'  # 明朝体
        style.font.size = Pt(11)  # フォントサイズを11ptに
        style.paragraph_format.line_spacing = 1.5  # 行間を1.5倍に
        style.paragraph_format.first_line_indent = Inches(0.25)  # 一字下げ
        style.paragraph_format.space_after = Pt(6)  # 段落後のスペース

# 見出しスタイルの設定
heading_configs = {
    'Heading 1': {'size': 18, 'bold': True, 'space_before': 24, 'space_after': 12},
    'Heading 2': {'size': 14, 'bold': True, 'space_before': 18, 'space_after': 6},
    'Heading 3': {'size': 12, 'bold': True, 'space_before': 12, 'space_after': 6},
}

for style_name, config in heading_configs.items():
    if style_name in doc.styles:
        style = doc.styles[style_name]
        style.font.name = 'MS Gothic'  # ゴシック体
        style.font.size = Pt(config['size'])
        style.font.bold = config['bold']
        style.paragraph_format.space_before = Pt(config['space_before'])
        style.paragraph_format.space_after = Pt(config['space_after'])
        style.paragraph_format.first_line_indent = Inches(0)  # 見出しは字下げなし

print("  ✅ スタイル設定完了")

# 段落に一字下げを適用
paragraph_count = 0
for para in doc.paragraphs:
    if para.style.name in ['Normal', 'Body Text', 'First Paragraph']:
        if para.paragraph_format.first_line_indent is None or para.paragraph_format.first_line_indent == 0:
            para.paragraph_format.first_line_indent = Inches(0.25)
            paragraph_count += 1

print(f"  ✅ {paragraph_count}個の段落に一字下げを適用")

# 保存
doc.save('temp_with_styles_v27.docx')
print("\n✅ スタイル改善完了: temp_with_styles_v27.docx")

# ステップ4: テーブルの挿入と改善
print("\nステップ4: テーブルの挿入と改善")
doc = Document('temp_with_styles_v27.docx')

table_inserted_count = 0
table_number = 1

for table_id, table_data in missing_tables.items():
    placeholder = f"【テーブル挿入位置:{table_id}】"
    
    for i, para in enumerate(doc.paragraphs):
        if placeholder in para.text:
            # プレースホルダーをクリア
            para.text = ""
            
            # 図表番号を追加
            caption_para = doc.paragraphs[i].insert_paragraph_before()
            caption_para.text = f"表{table_number}: {table_data.get('caption', table_id)}"
            caption_para.style = 'Caption'
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # テーブルを挿入
            rows = len(table_data['rows'])
            cols = len(table_data['headers'])
            table = doc.add_table(rows=rows+1, cols=cols)
            table.style = 'Table Grid'
            
            # ヘッダー行
            for j, header in enumerate(table_data['headers']):
                cell = table.rows[0].cells[j]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # データ行
            for row_idx, row_data in enumerate(table_data['rows']):
                for col_idx, cell_value in enumerate(row_data):
                    table.rows[row_idx+1].cells[col_idx].text = cell_value
            
            # テーブルの枠線を2.0ptに設定
            tbl = table._element
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            
            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '16')  # 2.0pt (8 * 2)
                border.set(qn('w:color'), '000000')
                tblBorders.append(border)
            
            tblPr.append(tblBorders)
            
            # テーブルを段落の後に移動
            para._element.addnext(table._element)
            
            table_inserted_count += 1
            table_number += 1
            print(f"  ✅ {table_id} を挿入（表{table_number-1}）")
            break

print(f"\n✅ {table_inserted_count}個のテーブルを挿入")

# 保存
doc.save('temp_with_tables_v27.docx')
print("\n✅ テーブル挿入完了: temp_with_tables_v27.docx")

# ステップ5: ページブレークを削除
print("\nステップ5: ページブレークを削除")
doc = Document('temp_with_tables_v27.docx')

page_break_count = 0
for para in doc.paragraphs:
    for run in para.runs:
        if 'w:br' in run._element.xml and 'w:type="page"' in run._element.xml:
            run._element.getparent().remove(run._element)
            page_break_count += 1

print(f"  ✅ {page_break_count}個のページブレークを削除")

# 保存
doc.save('complete_manuscript_v27.docx')
print("\n✅ ページブレーク削除完了: complete_manuscript_v27.docx")

# ステップ6: ヘッダー・フッターの追加
print("\nステップ6: ヘッダー・フッターの追加")
doc = Document('complete_manuscript_v27.docx')

# ヘッダーを設定
for section in doc.sections:
    header = section.header
    header.is_linked_to_previous = False
    
    # ヘッダーをクリア
    for para in header.paragraphs:
        para.clear()
    
    # ヘッダーテーブルを作成
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    header_table.autofit = False
    
    # 左セル：書籍タイトル
    left_cell = header_table.rows[0].cells[0]
    left_para = left_cell.paragraphs[0]
    left_para.text = "AIツアーコンダクターと歩く ベトナム・タイ周遊記"
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_run = left_para.runs[0]
    left_run.font.size = Pt(9)
    left_run.font.italic = True
    
    # 右セル：ページ番号
    right_cell = header_table.rows[0].cells[1]
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # ページ番号フィールドを追加
    run = right_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    run._element.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar2)
    
    run.font.size = Pt(11)
    run.font.bold = True
    
    # 右セルの背景色を設定
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '7CB8C5')
    right_cell._element.get_or_add_tcPr().append(shading_elm)
    
    # テーブルの枠線を削除
    tbl = header_table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)

# 表紙ページのヘッダー・フッターを非表示
doc.sections[0].different_first_page_header_footer = True

print("  ✅ ヘッダー・フッター設定完了")

# 最終保存
doc.save('complete_manuscript_v27_final.docx')

print("\n" + "="*60)
print("✅ 改善版v27作成完了！")
print("="*60)
print(f"\n出力ファイル: complete_manuscript_v27_final.docx")
print(f"総段落数: {len(doc.paragraphs)}")
print(f"総テーブル数: {len(doc.tables)}")
print(f"総セクション数: {len(doc.sections)}")
