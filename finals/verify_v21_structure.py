#!/usr/bin/env python3.11
# -*- coding: utf-8 -*-
"""
v21のWord文書の見出し構造を検証するスクリプト
"""

from docx import Document

def verify_word_structure(filename):
    """
    Word文書の見出し構造を検証
    
    Args:
        filename: 検証するWord文書のファイル名
    """
    doc = Document(filename)
    
    heading1_count = 0
    heading2_count = 0
    heading3_count = 0
    table_count = 0
    
    print(f"\n{'='*80}")
    print(f"ファイル: {filename}")
    print(f"{'='*80}\n")
    
    print("【見出し1の一覧】")
    for para in doc.paragraphs:
        if para.style.name == 'Heading 1':
            heading1_count += 1
            print(f"{heading1_count}. {para.text}")
        elif para.style.name == 'Heading 2':
            heading2_count += 1
        elif para.style.name == 'Heading 3':
            heading3_count += 1
    
    # テーブルの数を数える
    table_count = len(doc.tables)
    
    print(f"\n{'='*80}")
    print(f"【統計情報】")
    print(f"{'='*80}")
    print(f"見出し1の数: {heading1_count}")
    print(f"見出し2の数: {heading2_count}")
    print(f"見出し3の数: {heading3_count}")
    print(f"テーブルの数: {table_count}")
    print(f"{'='*80}\n")
    
    # 期待値との比較
    expected_heading1 = 17
    if heading1_count == expected_heading1:
        print(f"✅ 見出し1の数が正しいです（期待値: {expected_heading1}）")
    else:
        print(f"❌ 見出し1の数が期待値と異なります（期待値: {expected_heading1}, 実際: {heading1_count}）")
    
    if table_count > 0:
        print(f"✅ テーブルが{table_count}個含まれています")
    else:
        print(f"❌ テーブルが含まれていません")

if __name__ == "__main__":
    verify_word_structure("complete_manuscript_v21_final.docx")
