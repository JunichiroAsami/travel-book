#!/usr/bin/env python3.11
# -*- coding: utf-8 -*-
"""
Word文書に目次とページ番号を追加するスクリプト
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_page_number(section):
    """
    セクションのフッターにページ番号を追加
    
    Args:
        section: Word文書のセクション
    """
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ページ番号フィールドを追加
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def add_toc(doc):
    """
    文書の先頭に目次を追加
    
    Args:
        doc: Word文書
    """
    # 目次の見出しを追加
    toc_heading = doc.add_paragraph()
    toc_heading.style = 'Heading 1'
    toc_heading.add_run('目次')
    
    # 目次フィールドを追加
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-2" \\h \\z \\u'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    
    # ページ区切りを追加
    doc.add_page_break()

def add_toc_and_page_numbers(input_file, output_file):
    """
    Word文書に目次とページ番号を追加
    
    Args:
        input_file: 入力ファイル名
        output_file: 出力ファイル名
    """
    print("=" * 80)
    print("Word文書に目次とページ番号を追加")
    print("=" * 80)
    print()
    
    # Word文書を読み込み
    doc = Document(input_file)
    print(f"✓ {input_file} を読み込みました")
    
    # 目次を先頭に追加
    # 既存の内容を保持するため、新しいドキュメントを作成
    new_doc = Document()
    
    # 目次を追加
    add_toc(new_doc)
    
    # 既存の内容をコピー
    for element in doc.element.body:
        new_doc.element.body.append(element)
    
    # ページ番号を追加
    for section in new_doc.sections:
        add_page_number(section)
    
    print("✓ 目次を追加しました")
    print("✓ ページ番号を追加しました")
    
    # 保存
    new_doc.save(output_file)
    print()
    print("=" * 80)
    print(f"✓ 完了: {output_file}")
    print("=" * 80)
    print()
    print("次のステップ:")
    print("  1. Wordで開いて目次を更新: 目次を右クリック → 「フィールドの更新」")
    print("  2. 最終確認: 誤字脱字、フォーマットの確認")
    print("  3. Amazon KDPに提出")

if __name__ == "__main__":
    input_file = "complete_manuscript_v21_final.docx"
    output_file = "complete_manuscript_v21_final_with_toc.docx"
    
    add_toc_and_page_numbers(input_file, output_file)
