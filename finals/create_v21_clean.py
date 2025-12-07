#!/usr/bin/env python3.11
"""
Pandocで変換したWord文書にv19のヘッダーとフッターを適用
"""

from docx import Document
from copy import deepcopy

print("=" * 80)
print("v21のWord文書を作成（クリーンな方法）")
print("=" * 80)
print()

# v19をテンプレートとして読み込み（ヘッダーとフッターのみ使用）
template = Document("complete_manuscript_v19_final.docx")
print("✓ v19をテンプレートとして読み込みました")

# v19のヘッダーとフッターを保存
header_xml = template.sections[0].header._element
footer_xml = template.sections[0].footer._element

print("✓ v19のヘッダーとフッターを保存しました")

# Pandocで変換したv21を読み込み
# まず、統合マークダウンをPandocで変換
import subprocess

print()
print("Pandocでv21の統合マークダウンを変換中...")

# 目次付きでWord変換
subprocess.run([
    "pandoc",
    "complete_manuscript_v21.md",
    "-o", "complete_manuscript_v21_temp.docx",
    "--reference-doc=complete_manuscript_v19_final.docx",
    "--toc",
    "--toc-depth=2"
], check=True)

print("✓ Pandocでの変換が完了しました")

# 変換したWord文書を読み込み
doc = Document("complete_manuscript_v21_temp.docx")
print(f"✓ 変換したWord文書を読み込みました（{len(doc.paragraphs)}段落）")

# v19のヘッダーとフッターを適用
print()
print("v19のヘッダーとフッターを適用中...")

for section in doc.sections:
    # ヘッダーをコピー
    section.header._element.clear()
    for child in header_xml:
        section.header._element.append(deepcopy(child))
    
    # フッターをコピー
    section.footer._element.clear()
    for child in footer_xml:
        section.footer._element.append(deepcopy(child))

print("✓ すべてのセクションにヘッダーとフッターを適用しました")

# 保存
output_file = "complete_manuscript_v21_final.docx"
doc.save(output_file)

print()
print("=" * 80)
print(f"✓ 完了: {output_file}")
print("=" * 80)

# 検証
heading1_count = sum(1 for para in doc.paragraphs if para.style.name == 'Heading 1')
heading2_count = sum(1 for para in doc.paragraphs if para.style.name == 'Heading 2')
heading3_count = sum(1 for para in doc.paragraphs if para.style.name == 'Heading 3')

print()
print("【検証結果】")
print(f"  見出し1の数: {heading1_count}")
print(f"  見出し2の数: {heading2_count}")
print(f"  見出し3の数: {heading3_count}")
print(f"  総段落数: {len(doc.paragraphs)}")
print(f"  セクション数: {len(doc.sections)}")
