#!/usr/bin/env python3.11
"""
v10から軽量版参照文書を作成

目的:
- Amazon KDP用ドキュメントに含めるサンプル参照文書を作成
- 最初の2章のみを含む（ファイルサイズを削減）
- スタイル情報は完全に保持
"""

from docx import Document

# v10を読み込み
doc = Document("complete_manuscript_v10.docx")

print("=" * 80)
print("v10から軽量版参照文書を作成")
print("=" * 80)
print()

# 元の文書の情報
print(f"元の文書の段落数: {len(doc.paragraphs):,}")
print(f"元の文書のセクション数: {len(doc.sections)}")
print()

# 最初の2章（見出し1が2つ）を保持
heading1_count = 0
keep_paragraphs = []

for para in doc.paragraphs:
    if para.style.name == 'Heading 1':
        heading1_count += 1
        if heading1_count > 2:
            break
    keep_paragraphs.append(para)

print(f"保持する段落数: {len(keep_paragraphs):,}")
print(f"保持する見出し1の数: {min(heading1_count, 2)}")
print()

# 新しい文書を作成
new_doc = Document()

# セクション設定をコピー
if len(doc.sections) > 0:
    section = new_doc.sections[0]
    source_section = doc.sections[0]
    
    # ページサイズ
    section.page_width = source_section.page_width
    section.page_height = source_section.page_height
    
    # 余白
    section.top_margin = source_section.top_margin
    section.bottom_margin = source_section.bottom_margin
    section.left_margin = source_section.left_margin
    section.right_margin = source_section.right_margin
    
    # ヘッダーとフッター
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    
    # ヘッダーをコピー
    for element in source_section.header._element:
        section.header._element.append(element)
    
    # フッターをコピー
    for element in source_section.footer._element:
        section.footer._element.append(element)

print("✓ セクション設定をコピーしました")

# スタイルをコピー
for style in doc.styles:
    if style.name not in new_doc.styles:
        try:
            new_doc.styles.add_style(style.name, style.type)
        except:
            pass

print("✓ スタイルをコピーしました")

# 段落をコピー
for para in keep_paragraphs:
    new_para = new_doc.add_paragraph()
    new_para.text = para.text
    new_para.style = para.style
    
    # 段落のフォーマットをコピー
    if para.paragraph_format:
        new_para.paragraph_format.alignment = para.paragraph_format.alignment
        new_para.paragraph_format.left_indent = para.paragraph_format.left_indent
        new_para.paragraph_format.right_indent = para.paragraph_format.right_indent
        new_para.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
        new_para.paragraph_format.space_before = para.paragraph_format.space_before
        new_para.paragraph_format.space_after = para.paragraph_format.space_after

print("✓ 段落をコピーしました")

# 保存
output_file = "../standards/amazon_kdp/reference_document_sample.docx"
new_doc.save(output_file)

print()
print("=" * 80)
print(f"✓ 完了: {output_file}")
print("=" * 80)
print()
print("【軽量版参照文書の情報】")
print(f"  段落数: {len(new_doc.paragraphs):,}")
print(f"  セクション数: {len(new_doc.sections)}")
print(f"  ページサイズ: {section.page_width.cm:.2f}cm × {section.page_height.cm:.2f}cm")
print(f"  余白: 上{section.top_margin.cm:.2f}cm、下{section.bottom_margin.cm:.2f}cm、左{section.left_margin.cm:.2f}cm、右{section.right_margin.cm:.2f}cm")
