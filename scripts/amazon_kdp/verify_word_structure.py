#!/usr/bin/env python3
"""
Word文書の見出し構造を確認

使用方法:
    cd /home/ubuntu/ai_travel_book_project/finals
    python3.11 scripts/verify_word_structure.py [ファイル名]

例:
    python3.11 scripts/verify_word_structure.py complete_manuscript_v18_final.docx

必要なライブラリ:
    pip3 install python-docx
"""

import sys
import os
from docx import Document

def verify_word_structure(docx_file):
    """Word文書の見出し構造を確認"""
    
    if not os.path.exists(docx_file):
        print(f'エラー: {docx_file} が見つかりません')
        return
    
    doc = Document(docx_file)
    
    heading1_list = []
    heading1_count = 0
    heading2_count = 0
    heading3_count = 0
    heading4_count = 0
    
    for para in doc.paragraphs:
        if para.style.name == 'Heading 1':
            heading1_list.append(para.text)
            heading1_count += 1
        elif para.style.name == 'Heading 2':
            heading2_count += 1
        elif para.style.name == 'Heading 3':
            heading3_count += 1
        elif para.style.name == 'Heading 4':
            heading4_count += 1
    
    print('=' * 80)
    print(f'Word文書の見出し構造: {docx_file}')
    print('=' * 80)
    print()
    print(f'段落数: {len(doc.paragraphs):,}')
    print(f'Heading 1: {heading1_count}')
    print(f'Heading 2: {heading2_count}')
    print(f'Heading 3: {heading3_count}')
    print(f'Heading 4: {heading4_count}')
    print()
    
    print('Heading 1の一覧:')
    print('-' * 80)
    for i, title in enumerate(heading1_list, start=1):
        print(f'{i:2}. {title}')
    print('-' * 80)
    print()
    
    # 検証（本書の場合: 17個のHeading 1が期待される）
    expected_heading1_count = 17
    print(f'期待されるHeading 1: {expected_heading1_count}')
    print(f'実際のHeading 1: {heading1_count}')
    print()
    
    if heading1_count == expected_heading1_count:
        print('✓ すべてのHeading 1が存在します')
    else:
        print(f'✗ Heading 1が{heading1_count}個です（期待値: {expected_heading1_count}）')
        if heading1_count < expected_heading1_count:
            print(f'   {expected_heading1_count - heading1_count}個のHeading 1が不足しています')
        else:
            print(f'   {heading1_count - expected_heading1_count}個のHeading 1が超過しています')
    
    print()
    
    # ファイルサイズを表示
    file_size = os.path.getsize(docx_file)
    print(f'ファイルサイズ: {file_size / 1024:.2f} KB ({file_size / (1024 * 1024):.2f} MB)')
    
    # Amazon KDPの推奨サイズ（50MB）と比較
    if file_size < 50 * 1024 * 1024:
        print('✓ ファイルサイズはAmazon KDPの推奨範囲内です（50MB以下）')
    else:
        print('✗ ファイルサイズがAmazon KDPの推奨範囲を超えています（50MB以上）')

if __name__ == "__main__":
    if len(sys.argv) < 2:
        # デフォルトのファイル名
        docx_file = 'complete_manuscript_v18_final.docx'
        print(f'ファイル名が指定されていないため、デフォルトのファイルを確認します: {docx_file}')
        print()
    else:
        docx_file = sys.argv[1]
    
    verify_word_structure(docx_file)
