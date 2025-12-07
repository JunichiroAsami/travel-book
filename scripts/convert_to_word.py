#!/usr/bin/env python3
"""
統合原稿v3をWord形式に変換するスクリプト
参考書籍のスタイルに準拠
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

def create_styled_document():
    """スタイル設定済みのDocumentを作成"""
    doc = Document()
    
    # ページ設定
    section = doc.sections[0]
    section.page_height = Inches(8.27)  # A4縦
    section.page_width = Inches(5.83)   # A5横相当
    section.top_margin = Inches(0.79)
    section.bottom_margin = Inches(0.79)
    section.left_margin = Inches(0.79)
    section.right_margin = Inches(0.79)
    
    # スタイル設定
    styles = doc.styles
    
    # 見出し1（章タイトル）
    heading1 = styles['Heading 1']
    heading1.font.name = 'Yu Gothic'
    heading1.font.size = Pt(18)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(0, 0, 0)
    
    # 見出し2（節）
    heading2 = styles['Heading 2']
    heading2.font.name = 'Yu Gothic'
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(0, 0, 0)
    
    # 見出し3（項）
    heading3 = styles['Heading 3']
    heading3.font.name = 'Yu Gothic'
    heading3.font.size = Pt(12)
    heading3.font.bold = True
    heading3.font.color.rgb = RGBColor(0, 0, 0)
    
    # 本文
    normal = styles['Normal']
    normal.font.name = 'Yu Mincho'
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    
    return doc

def parse_markdown_line(line):
    """Markdown行を解析してタイプと内容を返す"""
    # 見出し1（# ）
    if line.startswith('# ') and not line.startswith('## '):
        return 'heading1', line[2:].strip()
    # 見出し2（## ）
    elif line.startswith('## ') and not line.startswith('### '):
        return 'heading2', line[3:].strip()
    # 見出し3（### ）
    elif line.startswith('### ') and not line.startswith('#### '):
        return 'heading3', line[4:].strip()
    # 水平線
    elif line.strip() in ['---', '***', '___']:
        return 'hr', ''
    # 空行
    elif line.strip() == '':
        return 'empty', ''
    # 箇条書き
    elif line.strip().startswith('- ') or line.strip().startswith('* '):
        return 'list', line.strip()[2:]
    # 番号付きリスト
    elif re.match(r'^\d+\.\s', line.strip()):
        return 'numbered_list', re.sub(r'^\d+\.\s', '', line.strip())
    # 引用
    elif line.strip().startswith('> '):
        return 'quote', line.strip()[2:]
    # 通常のテキスト
    else:
        return 'text', line.strip()

def clean_markdown_formatting(text):
    """Markdown装飾を削除"""
    # 太字
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # 斜体
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    # コード
    text = re.sub(r'`(.+?)`', r'\1', text)
    # リンク
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text

def convert_markdown_to_word(input_file, output_file):
    """MarkdownファイルをWord形式に変換"""
    doc = create_styled_document()
    
    with open(input_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    skip_next_empty = False
    
    for line in lines:
        line_type, content = parse_markdown_line(line)
        
        if line_type == 'empty':
            if not skip_next_empty:
                # 空行は段落の区切りとして扱う
                pass
            skip_next_empty = False
            continue
        
        # Markdown装飾を削除
        content = clean_markdown_formatting(content)
        
        if line_type == 'heading1':
            # 章タイトルの前に改ページ（最初の章を除く）
            if len(doc.paragraphs) > 1:
                doc.add_page_break()
            p = doc.add_heading(content, level=1)
            skip_next_empty = True
            
        elif line_type == 'heading2':
            doc.add_heading(content, level=2)
            skip_next_empty = True
            
        elif line_type == 'heading3':
            doc.add_heading(content, level=3)
            skip_next_empty = True
            
        elif line_type == 'hr':
            # 水平線は段落区切りとして扱う
            p = doc.add_paragraph()
            p.add_run('　' * 10)  # 全角スペースで区切り
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        elif line_type == 'list':
            p = doc.add_paragraph(content, style='List Bullet')
            
        elif line_type == 'numbered_list':
            p = doc.add_paragraph(content, style='List Number')
            
        elif line_type == 'quote':
            p = doc.add_paragraph(content)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            
        elif line_type == 'text':
            if content:  # 空でない場合のみ追加
                doc.add_paragraph(content)
    
    # 目次を先頭に挿入
    doc.paragraphs[0].insert_paragraph_before('目次')
    
    doc.save(output_file)
    print(f"変換完了: {output_file}")

if __name__ == '__main__':
    input_file = '/home/ubuntu/ai_travel_book_project/finals/complete_manuscript_v3.md'
    output_file = '/home/ubuntu/ai_travel_book_project/finals/complete_manuscript_v3.docx'
    
    convert_markdown_to_word(input_file, output_file)
