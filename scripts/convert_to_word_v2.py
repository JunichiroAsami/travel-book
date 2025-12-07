#!/usr/bin/env python3
"""
統合原稿v3をWord形式に変換するスクリプト（サンプル準拠版）
参考：生成AIと政治・地域行政のスタイル
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def set_cell_background(cell, fill_color):
    """セルの背景色を設定"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_styled_document():
    """スタイル設定済みのDocumentを作成"""
    doc = Document()
    
    # ページ設定（A5サイズ）
    section = doc.sections[0]
    section.page_height = Cm(21.0)  # A5縦
    section.page_width = Cm(14.8)   # A5横
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    
    # ヘッダー設定（ページ番号）
    section.header_distance = Cm(1.0)
    
    # スタイル設定
    styles = doc.styles
    
    # 見出し1（章タイトル）- 緑色の背景ボックス
    heading1 = styles['Heading 1']
    heading1.font.name = 'Yu Gothic'
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(255, 255, 255)  # 白文字
    
    # 見出し2（節）- 黄色の背景ボックス
    heading2 = styles['Heading 2']
    heading2.font.name = 'Yu Gothic'
    heading2.font.size = Pt(13)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(0, 0, 0)
    
    # 見出し3（項）- 太字、下線
    heading3 = styles['Heading 3']
    heading3.font.name = 'Yu Gothic'
    heading3.font.size = Pt(11)
    heading3.font.bold = True
    heading3.font.underline = True
    heading3.font.color.rgb = RGBColor(0, 0, 0)
    
    # 本文
    normal = styles['Normal']
    normal.font.name = 'Yu Mincho'
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.6
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = Pt(10.5)  # 1文字分の字下げ
    
    return doc

def add_colored_heading(doc, text, level):
    """背景色付きの見出しを追加"""
    if level == 1:
        # 緑色の背景ボックス
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.right_indent = Cm(0)
        
        run = p.add_run(text)
        run.font.name = 'Yu Gothic'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        
        # 背景色を設定（緑色）
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '4A9B7F')  # 緑色
        p._element.get_or_add_pPr().append(shading_elm)
        
    elif level == 2:
        # 黄色の背景ボックス
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(10)
        
        run = p.add_run(text)
        run.font.name = 'Yu Gothic'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        # 背景色を設定（薄い黄色）
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'F4E5B8')  # 薄い黄色
        p._element.get_or_add_pPr().append(shading_elm)
        
    elif level == 3:
        p = doc.add_heading(text, level=3)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(6)

def parse_markdown_line(line):
    """Markdown行を解析してタイプと内容を返す"""
    if line.startswith('# ') and not line.startswith('## '):
        return 'heading1', line[2:].strip()
    elif line.startswith('## ') and not line.startswith('### '):
        return 'heading2', line[3:].strip()
    elif line.startswith('### ') and not line.startswith('#### '):
        return 'heading3', line[4:].strip()
    elif line.strip() in ['---', '***', '___']:
        return 'hr', ''
    elif line.strip() == '':
        return 'empty', ''
    elif line.strip().startswith('- ') or line.strip().startswith('* '):
        return 'list', line.strip()[2:]
    elif re.match(r'^\d+\.\s', line.strip()):
        return 'numbered_list', re.sub(r'^\d+\.\s', '', line.strip())
    elif line.strip().startswith('> '):
        return 'quote', line.strip()[2:]
    elif line.strip().startswith('**') and line.strip().endswith('**'):
        return 'bold', line.strip()[2:-2]
    else:
        return 'text', line.strip()

def clean_markdown_formatting(text):
    """Markdown装飾を削除"""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text

def convert_markdown_to_word(input_file, output_file):
    """MarkdownファイルをWord形式に変換"""
    doc = create_styled_document()
    
    with open(input_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    skip_next_empty = False
    in_list = False
    
    for i, line in enumerate(lines):
        line_type, content = parse_markdown_line(line)
        
        if line_type == 'empty':
            if in_list:
                in_list = False
            if not skip_next_empty:
                pass
            skip_next_empty = False
            continue
        
        content = clean_markdown_formatting(content)
        
        if line_type == 'heading1':
            if len(doc.paragraphs) > 1:
                doc.add_page_break()
            add_colored_heading(doc, content, 1)
            skip_next_empty = True
            in_list = False
            
        elif line_type == 'heading2':
            add_colored_heading(doc, content, 2)
            skip_next_empty = True
            in_list = False
            
        elif line_type == 'heading3':
            add_colored_heading(doc, content, 3)
            skip_next_empty = True
            in_list = False
            
        elif line_type == 'hr':
            p = doc.add_paragraph()
            p.add_run('　' * 10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            in_list = False
            
        elif line_type == 'list':
            p = doc.add_paragraph(content, style='List Bullet')
            p.paragraph_format.first_line_indent = Pt(0)  # リストは字下げしない
            in_list = True
            
        elif line_type == 'numbered_list':
            p = doc.add_paragraph(content, style='List Number')
            p.paragraph_format.first_line_indent = Pt(0)
            in_list = True
            
        elif line_type == 'quote':
            p = doc.add_paragraph(content)
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.right_indent = Cm(1.0)
            p.paragraph_format.first_line_indent = Pt(0)
            in_list = False
            
        elif line_type == 'bold':
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.bold = True
            in_list = False
            
        elif line_type == 'text':
            if content:
                doc.add_paragraph(content)
            in_list = False
    
    doc.save(output_file)
    print(f"変換完了: {output_file}")

if __name__ == '__main__':
    input_file = '/home/ubuntu/ai_travel_book_project/finals/complete_manuscript_v3.md'
    output_file = '/home/ubuntu/ai_travel_book_project/finals/complete_manuscript_v3_styled.docx'
    
    convert_markdown_to_word(input_file, output_file)
